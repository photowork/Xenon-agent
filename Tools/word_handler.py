import json
import time
import copy
import re
from typing import Dict, Any, List, Optional, Tuple
from pathlib import Path
import logging

try:
    from docx import Document
    from docx.shared import Inches, Pt, Cm, Emu, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
    from docx.enum.style import WD_STYLE_TYPE
    from docx.oxml.shared import OxmlElement, qn
    import docx
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    import mammoth
    MAMMOTH_AVAILABLE = True
except ImportError:
    MAMMOTH_AVAILABLE = False

try:
    import pandas as pd
    PANDAS_AVAILABLE = True
except ImportError:
    PANDAS_AVAILABLE = False

try:
    import win32com.client
    WIN32COM_AVAILABLE = True
except ImportError:
    WIN32COM_AVAILABLE = False
    logging.warning("pywin32 库不可用，Word 图片插入功能将受限。请尝试安装: pip install pywin32")


class WordHandler:
    # ── 学术预设常量 ──
    ACADEMIC_PRESETS = {
        "cn_academic": {
            "page_size": "A4",
            "margins_cm": {"top": 2.54, "bottom": 2.54, "left": 3.17, "right": 3.17},
            "body_font": "宋体",
            "body_font_size": 12.0,  # 小四
            "body_line_spacing": 1.5,
            "body_first_line_indent_cm": 0.74,  # 两个字符
            "heading_styles": {
                1: {"font": "黑体", "size": 16.0, "bold": True, "alignment": "center"},  # 三号
                2: {"font": "黑体", "size": 14.0, "bold": True, "alignment": "left"},  # 四号
                3: {"font": "黑体", "size": 13.0, "bold": True, "alignment": "left"},  # 小四
            },
            "label_figure": "图",
            "label_table": "表",
            "label_equation": "式",
        },
        "ieee_like": {
            "page_size": "A4",
            "margins_cm": {"top": 1.91, "bottom": 2.54, "left": 1.91, "right": 1.91},
            "body_font": "Times New Roman",
            "body_font_size": 10.0,
            "body_line_spacing": 1.0,
            "body_first_line_indent_cm": 0.0,
            "heading_styles": {
                1: {"font": "Times New Roman", "size": 14.0, "bold": True, "alignment": "center"},
                2: {"font": "Times New Roman", "size": 12.0, "bold": True, "alignment": "left"},
                3: {"font": "Times New Roman", "size": 11.0, "bold": True, "alignment": "left"},
            },
            "label_figure": "Figure",
            "label_table": "Table",
            "label_equation": "Equation",
        },
        "apa_like": {
            "page_size": "A4",
            "margins_cm": {"top": 2.54, "bottom": 2.54, "left": 2.54, "right": 2.54},
            "body_font": "Times New Roman",
            "body_font_size": 12.0,
            "body_line_spacing": 2.0,
            "body_first_line_indent_cm": 1.27,
            "heading_styles": {
                1: {"font": "Times New Roman", "size": 14.0, "bold": True, "alignment": "center"},
                2: {"font": "Times New Roman", "size": 12.0, "bold": True, "alignment": "left"},
                3: {"font": "Times New Roman", "size": 12.0, "bold": True, "italic": True, "alignment": "left"},
            },
            "label_figure": "Figure",
            "label_table": "Table",
            "label_equation": "Equation",
        },
    }

    # 论文占位章节模板
    THESIS_SECTIONS = {
        "zh": [
            ("摘要", 1), ("关键词", 0), ("引言", 1), ("相关工作", 1),
            ("方法", 1), ("实验", 1), ("结果", 1), ("讨论", 1), ("结论", 1), ("参考文献", 1),
        ],
        "en": [
            ("Abstract", 1), ("Keywords", 0), ("Introduction", 1), ("Related Work", 1),
            ("Method", 1), ("Experiments", 1), ("Results", 1), ("Discussion", 1), ("Conclusion", 1), ("References", 1),
        ],
    }

    ARTICLE_SECTIONS = {
        "zh": [
            ("摘要", 1), ("关键词", 0), ("引言", 1), ("方法", 1), ("结果与讨论", 1), ("结论", 1), ("参考文献", 1),
        ],
        "en": [
            ("Abstract", 1), ("Keywords", 0), ("Introduction", 1), ("Method", 1),
            ("Results and Discussion", 1), ("Conclusion", 1), ("References", 1),
        ],
    }

    def __init__(self, file_path: str = None):
        """
        初始化 Word 处理器
        :param file_path: Word 文件路径（可选，如果未提供则创建新文档）
        """
        self.file_path = file_path
        self.document = None
        # ── 学术辅助数据 ──
        self._caption_counters: Dict[str, int] = {}  # {"figure": 1, "table": 2, "equation": 0}
        self._reference_entries: Dict[str, Dict[str, Any]] = {}  # cite_key -> entry
        self._reference_order: List[str] = []  # 引用出现顺序
        self._review_counter: int = 0
        self._todo_counter: int = 0
        if not DOCX_AVAILABLE:
            logging.warning("python-docx 库不可用，Word 功能将受限。请尝试安装: pip install python-docx")
        if not PANDAS_AVAILABLE:
            logging.warning("pandas 未安装，部分功能可能受限。请安装: pip install pandas")
    
    def load_document(self, file_path: str = None) -> Dict[str, Any]:
        """
        加载现有Word文档
        :param file_path: Word 文件路径（可选，如果未提供则使用初始化时的路径）
        """
        import time
        start_time = time.time()
        
        if not DOCX_AVAILABLE:
            execution_time = time.time() - start_time
            return {
                "success": False, 
                "error": "python-docx 库不可用，无法加载Word文档",
                "execution_time": round(execution_time, 2)
            }
        
        try:
            path_to_load = file_path or self.file_path
            if path_to_load and Path(path_to_load).exists():
                self.document = Document(path_to_load)
                self.file_path = path_to_load
            else:
                # 创建新文档
                self.document = Document()
                self.file_path = path_to_load

            self._sync_caption_counters_from_document()
            
            execution_time = time.time() - start_time
            
            return {
                "success": True,
                "message": f"成功{'加载' if path_to_load and Path(path_to_load).exists() else '创建'}文档: {path_to_load or '新文档'}",
                "paragraph_count": len(self.document.paragraphs),
                "section_count": len(self.document.sections),
                "table_count": len(self.document.tables),
                "execution_time": round(execution_time, 2),
                "file_path": path_to_load or "新文档"
            }
        except NameError:
            execution_time = time.time() - start_time
            return {
                "success": False, 
                "error": "Document 类未定义，python-docx 库可能未正确安装",
                "execution_time": round(execution_time, 2)
            }
        except Exception as e:
            execution_time = time.time() - start_time
            return {
                "success": False, 
                "error": f"加载文档失败: {str(e)}",
                "execution_time": round(execution_time, 2)
            }

    def _iter_table_paragraphs(self, tables, seen_cells: Optional[set] = None):
        """递归遍历表格单元格段落，避免合并单元格被重复处理。"""
        if seen_cells is None:
            seen_cells = set()

        for table in tables:
            for row in table.rows:
                for cell in row.cells:
                    cell_id = id(cell._tc)
                    if cell_id in seen_cells:
                        continue
                    seen_cells.add(cell_id)
                    for paragraph in cell.paragraphs:
                        yield paragraph
                    yield from self._iter_table_paragraphs(cell.tables, seen_cells)

    def _iter_all_paragraphs(self):
        """遍历正文和表格中的所有段落。"""
        if not self.document:
            return
        for paragraph in self.document.paragraphs:
            yield paragraph
        yield from self._iter_table_paragraphs(self.document.tables)

    def _find_text_matches(self, text: str, old_text: str, replace_all: bool) -> List[int]:
        matches = []
        start = 0
        while True:
            index = text.find(old_text, start)
            if index == -1:
                break
            matches.append(index)
            if not replace_all:
                break
            start = index + len(old_text)
        return matches

    def _replace_text_in_paragraph(self, paragraph, old_text: str, new_text: str, replace_all: bool) -> int:
        """
        在段落文本节点内替换文本，尽量保留原有 run 样式、超链接/域等 XML 结构。
        跨文本节点命中时，新文本继承命中起点节点所在 run 的样式。
        """
        text_elements = paragraph._element.xpath('.//w:t')
        if not text_elements:
            return 0

        full_text = "".join(text_element.text or "" for text_element in text_elements)
        if old_text not in full_text:
            return 0

        matches = self._find_text_matches(full_text, old_text, replace_all)
        if not matches:
            return 0

        char_map = []
        for element_index, text_element in enumerate(text_elements):
            for char_index, _ in enumerate(text_element.text or ""):
                char_map.append((element_index, char_index))

        for start in reversed(matches):
            end = start + len(old_text)
            first_element_index, first_char_index = char_map[start]
            last_element_index, last_char_index = char_map[end - 1]

            if first_element_index == last_element_index:
                text_element = text_elements[first_element_index]
                original_text = text_element.text or ""
                text_element.text = original_text[:first_char_index] + new_text + original_text[last_char_index + 1:]
                continue

            first_element = text_elements[first_element_index]
            last_element = text_elements[last_element_index]
            first_text = first_element.text or ""
            last_text = last_element.text or ""
            first_element.text = first_text[:first_char_index] + new_text
            for element_index in range(first_element_index + 1, last_element_index):
                text_elements[element_index].text = ""
            last_element.text = last_text[last_char_index + 1:]

        return len(matches)

    def _caption_labels_by_type(self) -> Dict[str, List[str]]:
        labels = {
            "figure": ["图", "Figure"],
            "table": ["表", "Table"],
            "equation": ["式", "Equation"],
        }
        for preset in self.ACADEMIC_PRESETS.values():
            labels["figure"].append(preset.get("label_figure", ""))
            labels["table"].append(preset.get("label_table", ""))
            labels["equation"].append(preset.get("label_equation", ""))
        return {
            key: list(dict.fromkeys(label for label in value if label))
            for key, value in labels.items()
        }

    def _sync_caption_counters_from_document(self) -> None:
        """从已有题注文本中同步最大编号，避免重新加载后重复从 1 开始。"""
        if not self.document:
            return

        labels_by_type = self._caption_labels_by_type()
        counters = {key: self._caption_counters.get(key, 0) for key in labels_by_type}
        for paragraph in self._iter_all_paragraphs():
            text = paragraph.text.strip()
            if not text:
                continue
            for target_type, labels in labels_by_type.items():
                for label in labels:
                    match = re.match(
                        rf"^{re.escape(label)}\s*([0-9]+)(?:\s|[：:.\-、)]|$)",
                        text,
                        re.IGNORECASE,
                    )
                    if match:
                        counters[target_type] = max(counters[target_type], int(match.group(1)))
                        break

        self._caption_counters.update(counters)

    def _caption_sequence_name(self, target_type: str) -> str:
        return {
            "figure": "Figure",
            "table": "Table",
            "equation": "Equation",
        }.get(target_type, target_type)

    def _add_seq_field(self, paragraph, sequence_name: str, display_text: str):
        fld_char_begin = OxmlElement('w:fldChar')
        fld_char_begin.set(qn('w:fldCharType'), 'begin')
        fld_code = OxmlElement('w:instrText')
        fld_code.set(qn('xml:space'), 'preserve')
        fld_code.text = f' SEQ {sequence_name} \\* ARABIC '
        fld_char_separate = OxmlElement('w:fldChar')
        fld_char_separate.set(qn('w:fldCharType'), 'separate')
        fld_char_end = OxmlElement('w:fldChar')
        fld_char_end.set(qn('w:fldCharType'), 'end')

        paragraph.add_run()._element.append(fld_char_begin)
        paragraph.add_run()._element.append(fld_code)
        paragraph.add_run()._element.append(fld_char_separate)
        display_run = paragraph.add_run(display_text)
        paragraph.add_run()._element.append(fld_char_end)
        return display_run

    def create_new_document(self) -> Dict[str, Any]:
        """
        创建新的空白文档
        """
        if not DOCX_AVAILABLE:
            return {"success": False, "error": "python-docx 库不可用，无法创建新文档"}
        
        try:
            self.document = Document()
            self.file_path = None
            return {
                "success": True,
                "message": "成功创建新文档",
                "paragraph_count": 0,
                "section_count": 1
            }
        except Exception as e:
            return {"success": False, "error": f"创建新文档失败: {str(e)}"}
    
    def read_paragraphs(self, start: int = 0, end: Optional[int] = None) -> Dict[str, Any]:
        """
        读取文档段落内容
        :param start: 起始段落索引
        :param end: 结束段落索引（不包含）
        """
        if not DOCX_AVAILABLE:
            return {"success": False, "error": "python-docx 库不可用，无法读取段落"}
        
        try:
            if not self.document:
                return {"success": False, "error": "文档未加载"}
            
            paragraphs = self.document.paragraphs
            if end is None:
                end = len(paragraphs)
            
            # 确保索引在有效范围内
            start = max(0, min(start, len(paragraphs)))
            end = max(start, min(end, len(paragraphs)))
            
            content = []
            for i, para in enumerate(paragraphs[start:end], start):
                content.append({
                    "index": i,
                    "text": para.text
                })
            
            return {
                "success": True,
                "paragraphs": content,
                "total_count": len(content)
            }
        except Exception as e:
            return {"success": False, "error": f"读取段落失败: {str(e)}"}
    



    def save_document(self, output_path: Optional[str] = None) -> Dict[str, Any]:
        """
        保存文档
        :param output_path: 输出路径（可选，默认使用原路径）
        """
        if not DOCX_AVAILABLE:
            return {"success": False, "error": "python-docx 库不可用，无法保存文档"}
        
        try:
            if not self.document:
                return {"success": False, "error": "文档未加载"}
            
            save_path = output_path or self.file_path
            if not save_path:
                return {"success": False, "error": "未指定保存路径"}
            
            # 处理路径
            save_input_path = Path(save_path)
            if save_input_path.is_absolute():
                save_abs_path = save_input_path.resolve()
            else:
                save_abs_path = Path.cwd() / save_path
                save_abs_path = save_abs_path.resolve()
            
            # 获取保存前的文件大小（如果文件已存在）
            file_size_before = 0
            if save_abs_path.exists():
                file_size_before = save_abs_path.stat().st_size
            
            self.document.save(str(save_abs_path))
            
            # 获取保存后的文件大小
            file_size_after = save_abs_path.stat().st_size
            
            return {
                "success": True,
                "message": f"文档已保存到: {str(save_abs_path)}",
                "path": str(save_abs_path),
                "file_size_before": file_size_before,
                "file_size_after": file_size_after,
                "size_change": file_size_after - file_size_before
            }
        except Exception as e:
            return {"success": False, "error": f"保存文档失败: {str(e)}"}

    def add_paragraph(self, text: str, style: str = "Normal", position: int = -1) -> Dict[str, Any]:
        """
        在指定位置添加段落
        :param text: 要添加的文本内容
        :param style: 段落样式名称，如 'Normal'、'Heading 1' 等，默认为 'Normal'
        :param position: 插入位置索引，-1表示添加到文档末尾，默认为 -1
        :return: 包含操作结果信息的字典
        """
        if not DOCX_AVAILABLE:
            return {"success": False, "error": "python-docx 库不可用，无法添加段落"}
        
        try:
            if not self.document:
                return {"success": False, "error": "文档未加载"}
            
            if position < -1:
                return {"success": False, "error": f"插入位置 {position} 无效，必须大于等于 -1"}
            
            paragraph_count = len(self.document.paragraphs)
            
            if position == -1 or position >= paragraph_count:
                # 添加到末尾
                if style:
                    paragraph = self.document.add_paragraph(text, style=style)
                else:
                    paragraph = self.document.add_paragraph(text)
                paragraph_index = len(self.document.paragraphs) - 1
            else:
                # 使用 python-docx 原生插入能力，在目标段落前插入新段落
                paragraph = self.document.paragraphs[position].insert_paragraph_before(text)
                if style:
                    paragraph.style = style
                paragraph_index = position
            
            return {
                "success": True,
                "message": "段落添加成功",
                "paragraph_index": paragraph_index
            }
        except Exception as e:
            return {"success": False, "error": f"添加段落失败: {str(e)}"}

    def add_heading(self, text: str, level: int = 1, position: int = -1) -> Dict[str, Any]:
        """
        添加标题（1-9级）
        :param text: 标题文本内容
        :param level: 标题级别，范围为 1-9，数字越小级别越高，默认为 1
        :param position: 插入位置索引，-1表示添加到文档末尾，默认为 -1
        :return: 包含操作结果信息的字典
        """
        if not DOCX_AVAILABLE:
            return {"success": False, "error": "python-docx 库不可用，无法添加标题"}
        
        try:
            if not self.document:
                return {"success": False, "error": "文档未加载"}
            
            if level < 1 or level > 9:
                return {"success": False, "error": "标题级别必须在1-9之间"}
            
            # 由于docx库限制，我们简单地在末尾添加标题
            heading = self.document.add_heading(text, level=level)
            
            return {
                "success": True,
                "message": f"成功添加{level}级标题",
                "heading_level": level,
                "heading_text": text
            }
        except Exception as e:
            return {"success": False, "error": f"添加标题失败: {str(e)}"}

    def _move_table_to_position(self, table, position: int) -> None:
        if position is None or position == -1:
            return
        paragraphs = self.document.paragraphs
        if position >= len(paragraphs):
            return
        table_element = table._element
        table_element.getparent().remove(table_element)
        paragraphs[position]._element.addprevious(table_element)

    def replace_text(self, old_text: str, new_text: str, replace_all: bool = True) -> Dict[str, Any]:
        """
        替换文本内容
        :param old_text: 要被替换的原始文本
        :param new_text: 用于替换的新文本内容
        :param replace_all: 是否替换文档中所有匹配项，True替换全部，False仅替换第一个匹配项，默认为 True
        :return: 包含操作结果信息的字典，包括实际替换次数
        """
        if not DOCX_AVAILABLE:
            return {"success": False, "error": "python-docx 库不可用，无法替换文本"}
        
        try:
            if not self.document:
                return {"success": False, "error": "文档未加载"}

            if old_text is None or old_text == "":
                return {"success": False, "error": "要替换的文本不能为空"}
            if new_text is None:
                new_text = ""
            
            replacements = 0

            for paragraph in self.document.paragraphs:
                count = self._replace_text_in_paragraph(paragraph, old_text, new_text, replace_all)
                if count:
                    replacements += count
                    if not replace_all:
                        return {
                            "success": True,
                            "message": f"替换完成，共替换了 {replacements} 处文本",
                            "replacements": replacements
                        }
            
            for paragraph in self._iter_table_paragraphs(self.document.tables):
                count = self._replace_text_in_paragraph(paragraph, old_text, new_text, replace_all)
                if count:
                    replacements += count
                    if not replace_all:
                        return {
                            "success": True,
                            "message": f"替换完成，共替换了 {replacements} 处文本",
                            "replacements": replacements
                        }
            
            return {
                "success": True,
                "message": f"替换完成，共替换了 {replacements} 处文本",
                "replacements": replacements
            }
        except Exception as e:
            return {"success": False, "error": f"替换文本失败: {str(e)}"}

    def create_table(self, rows: int, cols: int, headers: Optional[List[str]] = None, data: Optional[List[List[str]]] = None, position: int = -1) -> Dict[str, Any]:
        """
        创建指定行列的表格
        :param rows: 表格行数，必须为正整数
        :param cols: 表格列数，必须为正整数
        :param headers: 表头列表，如果提供则第一行为表头，默认为 None
        :param data: 表格数据，二维列表格式 [[row1_col1, row1_col2], [row2_col1, row2_col2], ...]，默认为 None
        :param position: 插入位置索引，-1表示添加到文档末尾，默认为 -1
        :return: 包含操作结果信息的字典
        """
        try:
            if not self.document:
                return {"success": False, "error": "文档未加载"}
            
            if rows <= 0 or cols <= 0:
                return {"success": False, "error": "行数和列数必须大于0"}
            if position < -1:
                return {"success": False, "error": f"插入位置 {position} 无效，必须大于等于 -1"}
            
            table = self.document.add_table(rows=rows, cols=cols)
            table.style = 'Table Grid'
            
            # 如果提供了表头，设置表头
            row_offset = 0
            if headers and len(headers) == cols:
                header_cells = table.rows[0].cells
                for i, header in enumerate(headers):
                    header_cells[i].text = str(header)
                row_offset = 1
            
            # 如果提供了数据，填充表格
            if data:
                for row_idx, row_data in enumerate(data):
                    target_row_idx = row_idx + row_offset
                    if target_row_idx < rows:  # 确保不超过表格行数
                        row_cells = table.rows[target_row_idx].cells
                        for col_idx, cell_data in enumerate(row_data):
                            if col_idx < cols:  # 确保不超过表格列数
                                row_cells[col_idx].text = str(cell_data)

            self._move_table_to_position(table, position)
            
            return {
                "success": True,
                "message": f"成功创建 {rows}x{cols} 表格",
                "table_rows": rows,
                "table_cols": cols
            }
        except Exception as e:
            return {"success": False, "error": f"创建表格失败: {str(e)}"}

    def add_table_from_data(self, data_list: List[List[str]], headers: Optional[List[str]] = None, position: int = -1) -> Dict[str, Any]:
        """
        从二维数据列表直接创建表格
        :param data_list: 二维数据列表
        :param headers: 表头列表（可选）
        :param position: 插入位置，-1表示末尾
        """
        if not DOCX_AVAILABLE:
            return {"success": False, "error": "python-docx 库不可用，无法创建表格"}
        
        try:
            if not self.document:
                return {"success": False, "error": "文档未加载"}
            
            if not data_list:
                return {"success": False, "error": "数据列表不能为空"}
            if position < -1:
                return {"success": False, "error": f"插入位置 {position} 无效，必须大于等于 -1"}
            
            # 计算表格尺寸
            rows = len(data_list) + (1 if headers else 0)
            cols = max(len(row) for row in data_list) if data_list else 0
            if headers:
                cols = max(cols, len(headers))
            
            if rows <= 0 or cols <= 0:
                return {"success": False, "error": "数据无效"}
            
            table = self.document.add_table(rows=rows, cols=cols)
            table.style = 'Table Grid'
            
            row_offset = 0
            # 如果有表头，先设置表头
            if headers:
                header_cells = table.rows[0].cells
                for i, header in enumerate(headers):
                    if i < cols:
                        header_cells[i].text = str(header)
                row_offset = 1
            
            # 填充数据
            for row_idx, row_data in enumerate(data_list):
                row_cells = table.rows[row_idx + row_offset].cells
                for col_idx, cell_data in enumerate(row_data):
                    if col_idx < cols:
                        row_cells[col_idx].text = str(cell_data)

            self._move_table_to_position(table, position)
            
            return {
                "success": True,
                "message": f"成功从数据创建 {rows}x{cols} 表格",
                "table_rows": rows,
                "table_cols": cols
            }
        except Exception as e:
            return {"success": False, "error": f"从数据创建表格失败: {str(e)}"}

    def update_table_cell(self, table_index: int, row: int, col: int, value: str) -> Dict[str, Any]:
        """
        更新指定表格的单元格内容
        :param table_index: 表格索引
        :param row: 行索引
        :param col: 列索引
        :param value: 新的单元格值
        """
        if not DOCX_AVAILABLE:
            return {"success": False, "error": "python-docx 库不可用，无法更新表格单元格"}
        
        try:
            if not self.document:
                return {"success": False, "error": "文档未加载"}
            
            if table_index < 0 or table_index >= len(self.document.tables):
                return {"success": False, "error": f"表格索引 {table_index} 超出范围，文档中只有 {len(self.document.tables)} 个表格"}
            
            table = self.document.tables[table_index]
            
            if row < 0 or col < 0 or row >= len(table.rows) or col >= len(table.columns):
                return {"success": False, "error": f"行列索引超出范围，表格大小为 {len(table.rows)}x{len(table.columns)}"}
            
            table.cell(row, col).text = str(value)
            
            return {
                "success": True,
                "message": f"成功更新表格 {table_index} 的 ({row}, {col}) 单元格",
                "value": str(value)
            }
        except Exception as e:
            return {"success": False, "error": f"更新表格单元格失败: {str(e)}"}

    def get_table_data(self, table_index: int) -> Dict[str, Any]:
        """
        获取指定表格的所有数据
        :param table_index: 表格索引
        """
        if not DOCX_AVAILABLE:
            return {"success": False, "error": "python-docx 库不可用，无法获取表格数据"}
        
        try:
            if not self.document:
                return {"success": False, "error": "文档未加载"}
            
            if table_index < 0 or table_index >= len(self.document.tables):
                return {"success": False, "error": f"表格索引 {table_index} 超出范围，文档中只有 {len(self.document.tables)} 个表格"}
            
            table = self.document.tables[table_index]
            table_data = []
            
            for row in table.rows:
                row_data = []
                for cell in row.cells:
                    row_data.append(cell.text)
                table_data.append(row_data)
            
            return {
                "success": True,
                "table_data": table_data,
                "rows": len(table_data),
                "cols": len(table_data[0]) if table_data else 0
            }
        except Exception as e:
            return {"success": False, "error": f"获取表格数据失败: {str(e)}"}

    def set_paragraph_format(self, paragraph_index: int, alignment: str = None, spacing_before: float = None, spacing_after: float = None) -> Dict[str, Any]:
        """
        设置段落格式
        :param paragraph_index: 段落索引，从0开始计数
        :param alignment: 对齐方式，可选值：'left'(左对齐)、'center'(居中)、'right'(右对齐)、'justify'(两端对齐)，默认为 None（不改变对齐方式）
        :param spacing_before: 段前间距（磅），默认为 None（不改变段前间距）
        :param spacing_after: 段后间距（磅），默认为 None（不改变段后间距）
        :return: 包含操作结果信息的字典
        """
        try:
            if not self.document:
                return {"success": False, "error": "文档未加载"}
            
            if paragraph_index < 0 or paragraph_index >= len(self.document.paragraphs):
                return {"success": False, "error": f"段落索引 {paragraph_index} 超出范围，文档中只有 {len(self.document.paragraphs)} 个段落"}
            
            paragraph = self.document.paragraphs[paragraph_index]
            p_format = paragraph.paragraph_format
            
            if alignment:
                if alignment.lower() == 'left':
                    p_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
                elif alignment.lower() == 'center':
                    p_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
                elif alignment.lower() == 'right':
                    p_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                elif alignment.lower() == 'justify':
                    p_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            
            if spacing_before is not None:
                p_format.space_before = Pt(spacing_before)
            
            if spacing_after is not None:
                p_format.space_after = Pt(spacing_after)
            
            return {
                "success": True,
                "message": f"成功设置段落 {paragraph_index} 格式"
            }
        except Exception as e:
            return {"success": False, "error": f"设置段落格式失败: {str(e)}"}

    def set_font_style(self, text_range: str, bold: bool = False, italic: bool = False, underline: bool = False, font_size: float = None, font_name: str = None) -> Dict[str, Any]:
        """
        设置字体样式（目前简单实现为对整个文档中匹配文本的样式设置）
        :param text_range: 文本范围或要设置样式的具体文本内容
        :param bold: 是否加粗，True为加粗，False为不加粗，默认为 False
        :param italic: 是否斜体，True为斜体，False为不斜体，默认为 False
        :param underline: 是否添加下划线，True为添加，False为不添加，默认为 False
        :param font_size: 字体大小（磅），如 12.0 表示 12 磅，默认为 None（不改变字体大小）
        :param font_name: 字体名称，如 '宋体'、'Arial' 等，默认为 None（不改变字体）
        :return: 包含操作结果信息的字典
        """
        if not DOCX_AVAILABLE:
            return {"success": False, "error": "python-docx 库不可用，无法设置字体样式"}
        
        try:
            if not self.document:
                return {"success": False, "error": "文档未加载"}
            
            # 遍历所有段落，对匹配的文本设置样式
            modifications = 0
            for paragraph in self.document.paragraphs:
                for run in paragraph.runs:
                    if text_range in run.text:
                        if bold:
                            run.bold = bold
                        if italic:
                            run.italic = italic
                        if underline:
                            run.underline = underline
                        if font_size:
                            run.font.size = Pt(font_size)
                        if font_name:
                            run.font.name = font_name
                        modifications += 1
            
            # 也处理表格中的文本
            for table in self.document.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                if text_range in run.text:
                                    if bold:
                                        run.bold = bold
                                    if italic:
                                        run.italic = italic
                                    if underline:
                                        run.underline = underline
                                    if font_size:
                                        run.font.size = Pt(font_size)
                                    if font_name:
                                        run.font.name = font_name
                                    modifications += 1
            
            return {
                "success": True,
                "message": f"成功设置文本样式，修改了 {modifications} 处文本",
                "modifications": modifications
            }
        except Exception as e:
            return {"success": False, "error": f"设置字体样式失败: {str(e)}"}

    def insert_page_break(self, position: int = -1) -> Dict[str, Any]:
        """
        插入分页符
        :param position: 插入位置，-1表示末尾
        """
        if not DOCX_AVAILABLE:
            return {"success": False, "error": "python-docx 库不可用，无法添加分页符"}
        
        try:
            if not self.document:
                return {"success": False, "error": "文档未加载"}
            
            # 在末尾添加分页符
            self.document.add_page_break()
            
            return {
                "success": True,
                "message": "分页符插入成功"
            }
        except Exception as e:
            return {"success": False, "error": f"插入分页符失败: {str(e)}"}

    def add_section_break(self, position: int = -1) -> Dict[str, Any]:
        """
        插入分节符
        :param position: 插入位置，-1表示末尾
        """
        if not DOCX_AVAILABLE:
            return {"success": False, "error": "python-docx 库不可用，无法添加分节符"}
        
        try:
            if not self.document:
                return {"success": False, "error": "文档未加载"}
            
            # 获取当前节
            sections = self.document.sections
            new_section = self.document.add_section()
            
            return {
                "success": True,
                "message": "分节符插入成功",
                "section_count": len(self.document.sections)
            }
        except Exception as e:
            return {"success": False, "error": f"插入分节符失败: {str(e)}"}

    def get_document_structure(self) -> Dict[str, Any]:
        """
        获取文档结构信息（段落数、表格数等）
        """
        try:
            if not self.document:
                return {"success": False, "error": "文档未加载"}
            
            return {
                "success": True,
                "paragraph_count": len(self.document.paragraphs),
                "table_count": len(self.document.tables),
                "section_count": len(self.document.sections),
                "page_count": "无法直接获取（需要保存后通过其他方式计算）"
            }
        except Exception as e:
            return {"success": False, "error": f"获取文档结构失败: {str(e)}"}

    def import_excel_to_tables(self, excel_file_path: str, sheet_names: Optional[List[str]] = None) -> Dict[str, Any]:
        """
        将Excel工作表直接转换为Word表格
        :param excel_file_path: Excel文件路径
        :param sheet_names: 工作表名称列表（可选，如果未提供则导入所有工作表）
        """
        if not DOCX_AVAILABLE:
            return {"success": False, "error": "python-docx 库不可用，无法导入Excel到表格"}
        
        try:
            if not self.document:
                return {"success": False, "error": "文档未加载"}
            
            if not PANDAS_AVAILABLE:
                return {"success": False, "error": "需要安装pandas库: pip install pandas"}
            
            if not Path(excel_file_path).exists():
                return {"success": False, "error": f"Excel文件不存在: {excel_file_path}"}
            
            try:
                # 读取Excel文件
                if sheet_names:
                    excel_data = pd.read_excel(excel_file_path, sheet_name=sheet_names)
                else:
                    excel_data = pd.read_excel(excel_file_path, sheet_name=None)
                
                imported_sheets = []
                
                # 如果是单个工作表（字典中只有一个键）
                if not isinstance(excel_data, dict):
                    sheet_name = "Sheet1"
                    df = excel_data
                    self._add_dataframe_as_table(df, sheet_name)
                    imported_sheets.append(sheet_name)
                else:
                    # 多个工作表
                    for sheet_name, df in excel_data.items():
                        self._add_dataframe_as_table(df, sheet_name)
                        imported_sheets.append(sheet_name)
                
                return {
                    "success": True,
                    "message": f"成功从Excel导入 {len(imported_sheets)} 个工作表为表格",
                    "imported_sheets": imported_sheets
                }
            except Exception as e:
                return {"success": False, "error": f"读取Excel文件失败: {str(e)}"}
        except Exception as e:
            return {"success": False, "error": f"导入Excel到表格失败: {str(e)}"}

    def _add_dataframe_as_table(self, df, title: str = None):
        """
        将DataFrame添加为Word表格
        :param df: pandas DataFrame
        :param title: 表格标题（可选）
        """
        if not DOCX_AVAILABLE:
            raise Exception("python-docx 库不可用，无法添加表格")
        
        if title:
            self.document.add_heading(title, level=2)
        
        # 创建表格（行数+1为了表头，列数为DataFrame的列数）
        table = self.document.add_table(rows=1, cols=len(df.columns))
        table.style = 'Table Grid'
        
        # 添加表头
        hdr_cells = table.rows[0].cells
        for i, column_name in enumerate(df.columns):
            hdr_cells[i].text = str(column_name)
        
        # 添加数据行
        for _, row in df.iterrows():
            row_cells = table.add_row().cells
            for i, value in enumerate(row):
                row_cells[i].text = str(value) if pd.notna(value) else ""

    def export_tables_to_excel(self, output_path: str) -> Dict[str, Any]:
        """
        将文档中的表格导出到Excel
        :param output_path: 输出Excel文件路径
        """
        if not DOCX_AVAILABLE:
            return {"success": False, "error": "python-docx 库不可用，无法导出表格到Excel"}
        
        try:
            if not self.document:
                return {"success": False, "error": "文档未加载"}
            
            if not PANDAS_AVAILABLE:
                return {"success": False, "error": "需要安装pandas库: pip install pandas"}
            
            if not output_path.endswith(('.xlsx', '.xls')):
                output_path += '.xlsx'
            
            try:
                tables_data = {}
                for i, table in enumerate(self.document.tables):
                    # 将表格转换为DataFrame
                    data = []
                    for row in table.rows:
                        row_data = []
                        for cell in row.cells:
                            row_data.append(cell.text)
                        data.append(row_data)
                    
                    if data:  # 确保表格不为空
                        df = pd.DataFrame(data[1:], columns=data[0])  # 第一行作为列名
                        tables_data[f'Table_{i+1}'] = df
                
                if tables_data:
                    # 写入Excel文件
                    with pd.ExcelWriter(output_path, engine='openpyxl') as writer:
                        for sheet_name, df in tables_data.items():
                            df.to_excel(writer, sheet_name=sheet_name, index=False)
                else:
                    return {"success": False, "error": "文档中没有找到任何表格"}
                
                return {
                    "success": True,
                    "message": f"成功将 {len(tables_data)} 个表格导出到Excel: {output_path}",
                    "output_path": output_path
                }
            except Exception as e:
                return {"success": False, "error": f"导出Excel失败: {str(e)}"}
        except Exception as e:
            return {"success": False, "error": f"导出表格到Excel失败: {str(e)}"}

    def merge_table_cells(self, table_index: int, top_row: int, left_col: int, bottom_row: int, right_col: int) -> Dict[str, Any]:
        """
        合并表格中的单元格
        :param table_index: 表格索引
        :param top_row: 顶部行索引
        :param left_col: 左侧列索引
        :param bottom_row: 底部行索引
        :param right_col: 右侧列索引
        """
        if not DOCX_AVAILABLE:
            return {"success": False, "error": "python-docx 库不可用，无法合并表格单元格"}
        
        try:
            if not self.document:
                return {"success": False, "error": "文档未加载"}
            
            if table_index < 0 or table_index >= len(self.document.tables):
                return {"success": False, "error": f"表格索引 {table_index} 超出范围，文档中只有 {len(self.document.tables)} 个表格"}
            
            table = self.document.tables[table_index]
            
            if top_row < 0 or bottom_row < 0 or top_row >= len(table.rows) or bottom_row >= len(table.rows) or top_row > bottom_row:
                return {"success": False, "error": f"行索引超出范围或不合法，表格有 {len(table.rows)} 行"}
            
            if left_col < 0 or right_col < 0 or left_col >= len(table.columns) or right_col >= len(table.columns) or left_col > right_col:
                return {"success": False, "error": f"列索引超出范围或不合法，表格有 {len(table.columns)} 列"}
            
            # 获取要合并的单元格
            top_left_cell = table.cell(top_row, left_col)
            bottom_right_cell = table.cell(bottom_row, right_col)
            
            # 执行合并操作
            merged_cell = top_left_cell.merge(bottom_right_cell)
            
            return {
                "success": True,
                "message": f"成功合并表格 {table_index} 中从({top_row},{left_col})到({bottom_row},{right_col})的单元格",
                "merged_cell": f"({top_row},{left_col})-({bottom_row},{right_col})"
            }
        except Exception as e:
            return {"success": False, "error": f"合并单元格失败: {str(e)}"}

    def split_table_cell(self, table_index: int, row: int, col: int, v_merge: bool = True) -> Dict[str, Any]:
        """
        拆分表格中的已合并单元格。
        注意：python-docx 目前不提供可靠的取消合并能力，本方法当前只会做安全检查并返回明确错误，
        不会尝试直接修改底层 XML 以免损坏文档。
        :param table_index: 表格索引
        :param row: 行索引
        :param col: 列索引
        :param v_merge: 保留参数，当前不生效
        """
        if not DOCX_AVAILABLE:
            return {"success": False, "error": "python-docx 库不可用，无法拆分表格单元格"}
        
        try:
            if not self.document:
                return {"success": False, "error": "文档未加载"}
            
            if table_index < 0 or table_index >= len(self.document.tables):
                return {"success": False, "error": f"表格索引 {table_index} 超出范围，文档中只有 {len(self.document.tables)} 个表格"}
            
            table = self.document.tables[table_index]
            
            if row < 0 or col < 0 or row >= len(table.rows) or col >= len(table.columns):
                return {"success": False, "error": f"行列索引超出范围，表格大小为 {len(table.rows)}x{len(table.columns)}"}
            
            cell = table.cell(row, col)
            
            # python-docx 目前不提供可靠的“取消合并单元格”能力，直接改 XML 容易损坏文档。
            tc_pr = cell._tc.tcPr
            has_vmerge = tc_pr is not None and getattr(tc_pr, "vMerge", None) is not None
            grid_span = getattr(tc_pr, "gridSpan", None) if tc_pr is not None else None
            grid_span_val = getattr(grid_span, "val", None)
            has_gridspan = grid_span_val not in (None, 1, "1")
            
            if has_vmerge or has_gridspan:
                return {
                    "success": False,
                    "error": "python-docx 当前无法可靠拆分已合并单元格。请重新创建表格，或改用 Word COM 自动化处理。"
                }
            
            return {
                "success": False,
                "message": f"单元格({row},{col})未被合并，无需拆分",
                "cell_position": f"({row},{col})"
            }
        except Exception as e:
            return {"success": False, "error": f"拆分单元格失败: {str(e)}"}

    def adjust_cell_width_height(self, table_index: int, row: int, col: int, width_inches: float = None, height_inches: float = None) -> Dict[str, Any]:
        """
        调整表格单元格的宽度和高度
        :param table_index: 表格索引
        :param row: 行索引
        :param col: 列索引
        :param width_inches: 宽度（英寸）
        :param height_inches: 高度（英寸）
        """
        if not DOCX_AVAILABLE:
            return {"success": False, "error": "python-docx 库不可用，无法调整单元格宽高"}
        
        try:
            if not self.document:
                return {"success": False, "error": "文档未加载"}
            
            if table_index < 0 or table_index >= len(self.document.tables):
                return {"success": False, "error": f"表格索引 {table_index} 超出范围，文档中只有 {len(self.document.tables)} 个表格"}
            
            table = self.document.tables[table_index]
            
            if row < 0 or col < 0 or row >= len(table.rows) or col >= len(table.columns):
                return {"success": False, "error": f"行列索引超出范围，表格大小为 {len(table.rows)}x{len(table.columns)}"}
            
            cell = table.cell(row, col)
            
            # 调整单元格宽度
            if width_inches is not None:
                cell.width = Inches(width_inches)
            
            # 调整单元格高度（需要调整所在行的高度）
            if height_inches is not None:
                from docx.oxml.shared import OxmlElement, qn
                tr = table.rows[row]._tr
                trPr = tr.get_or_add_trPr()
                trHeight = trPr.xpath('./w:trHeight')[0] if trPr.xpath('./w:trHeight') else None
                
                if trHeight is None:
                    trHeight = OxmlElement('w:trHeight')
                    trPr.append(trHeight)
                
                # 设置行高（单位为twips，1英寸=1440 twips）
                trHeight.set(qn('w:val'), str(int(height_inches * 1440)))
                trHeight.set(qn('w:unit'), 'dxa')  # 设置单位为dxa (twips)
            
            return {
                "success": True,
                "message": f"成功调整表格 {table_index} 中({row},{col})单元格的尺寸",
                "width_inches": width_inches,
                "height_inches": height_inches
            }
        except Exception as e:
            return {"success": False, "error": f"调整单元格尺寸失败: {str(e)}"}
    
    def _resolve_path(self, file_path: str) -> Path:
        """将输入路径解析为绝对路径。"""
        input_path = Path(file_path)
        if input_path.is_absolute():
            return input_path.resolve()
        return (Path.cwd() / input_path).resolve()
    
    def _insert_inline_picture_with_docx(
        self,
        word_abs_path: Path,
        image_abs_path: Path,
        width: Optional[float] = None,
        height: Optional[float] = None,
        target_text: Optional[str] = None
    ) -> Dict[str, Any]:
        """
        使用 python-docx 进行跨平台的嵌入型图片插入。
        仅支持嵌入型图片，不支持文字环绕和浮动定位。
        """
        if not DOCX_AVAILABLE:
            return {"success": False, "error": "python-docx 库不可用，无法插入嵌入型图片"}
        
        try:
            file_size_before = word_abs_path.stat().st_size
            document = Document(str(word_abs_path))
            target_paragraph = None
            inserted_paragraph_index = None
            
            if target_text:
                for paragraph_index, paragraph in enumerate(document.paragraphs):
                    if target_text in paragraph.text:
                        target_paragraph = paragraph
                        inserted_paragraph_index = paragraph_index
                        break
                
                if target_paragraph is None:
                    return {"success": False, "error": f"未找到目标文本: {target_text}"}
                
                paragraph = target_paragraph.insert_paragraph_before("")
            else:
                paragraph = document.add_paragraph()
                inserted_paragraph_index = len(document.paragraphs) - 1
            
            run = paragraph.add_run()
            picture_kwargs = {}
            if width is not None:
                picture_kwargs["width"] = Pt(width)
            if height is not None:
                picture_kwargs["height"] = Pt(height)
            
            run.add_picture(str(image_abs_path), **picture_kwargs)
            document.save(str(word_abs_path))
            
            file_size_after = word_abs_path.stat().st_size
            if file_size_after <= file_size_before:
                return {
                    "success": False,
                    "error": f"图片插入失败：文件大小未增加（插入前: {file_size_before} 字节，插入后: {file_size_after} 字节）"
                }
            
            return {
                "success": True,
                "message": "成功插入图片（嵌入型，跨平台模式）",
                "wrap_type": "嵌入型",
                "output_file": str(word_abs_path),
                "file_size_before": file_size_before,
                "file_size_after": file_size_after,
                "size_increase": file_size_after - file_size_before,
                "backend": "python-docx",
                "inserted_paragraph_index": inserted_paragraph_index,
                "note": "当前模式不支持浮动定位和文字环绕；如需这些能力，请在 Windows + Word + pywin32 环境下使用。"
            }
        except Exception as e:
            return {"success": False, "error": f"插入嵌入型图片失败: {str(e)}"}

    def insert_picture_with_wrap(self, word_path: str, image_path: str, wrap_type: str = "嵌入型", 
                                   position_x: Optional[float] = None, position_y: Optional[float] = None,
                                   width: Optional[float] = None, height: Optional[float] = None,
                                   target_text: Optional[str] = None) -> Dict[str, Any]:
        """
        在Word文档中插入图片并设置文字环绕类型。
        兼容性说明：
        1. "嵌入型"(I) 使用 python-docx 实现，支持 Windows / Linux / macOS。
        2. 其他环绕类型（如四周型、紧密型、浮于文字上方等）依赖 Microsoft Word COM 自动化，
           仅支持 Windows + 已安装 Word + pywin32 的环境。
        3. 在非 Windows 环境中，请优先使用 "嵌入型"(I)。
        
        :param word_path: Word文档路径
        :param image_path: 图片路径
        :param wrap_type: 文字环绕类型，可选值：
                         - "四周型"(Q) - wdWrapSquare = 0
                         - "紧密型"(T) - wdWrapTight = 1
                         - "穿越型"(H) - wdWrapThrough = 2
                         - "上下型"(O) - wdWrapTopBottom = 3
                         - "衬于文字下方"(B) - wdWrapBehind = 5
                         - "浮于文字上方"(F) - wdWrapFront = 4
                         - "嵌入型"(I) - wdWrapInline = 7
        :param position_x: 图片水平位置（磅，1英寸=72磅），仅对 Windows 下的非嵌入型有效
        :param position_y: 图片垂直位置（磅，1英寸=72磅），仅对 Windows 下的非嵌入型有效
        :param width: 图片宽度（磅），None表示保持原始比例
        :param height: 图片高度（磅），None表示保持原始比例
        :param target_text: 目标文本位置，如果提供则在该文本处插入图片；跨平台模式下仅支持插入嵌入型图片
        :return: 包含成功状态和结果/错误的字典
        """
        word = None
        doc = None
        
        try:
            word_abs_path = self._resolve_path(word_path)
            image_abs_path = self._resolve_path(image_path)
            
            if not word_abs_path.exists():
                return {"success": False, "error": f"Word文档不存在: {word_path}"}
            
            if not image_abs_path.exists():
                return {"success": False, "error": f"图片文件不存在: {image_path}"}
            
            # 获取插入前的文件大小
            file_size_before = word_abs_path.stat().st_size
            
            wrap_type_map = {
                "四周型": 0,
                "Q": 0,
                "紧密型": 1,
                "T": 1,
                "穿越型": 2,
                "H": 2,
                "上下型": 3,
                "O": 3,
                "浮于文字上方": 4,
                "F": 4,
                "衬于文字下方": 5,
                "B": 5,
                "嵌入型": 7,
                "I": 7
            }
            
            if wrap_type not in wrap_type_map:
                return {"success": False, "error": f"无效的文字环绕类型: {wrap_type}。可选值: 四周型(Q), 紧密型(T), 穿越型(H), 上下型(O), 衬于文字下方(B), 浮于文字上方(F), 嵌入型(I)"}
            
            wrap_value = wrap_type_map[wrap_type]
            
            # 嵌入型图片可以直接通过 python-docx 实现，从而支持 Linux/macOS。
            if wrap_value == 7:
                return self._insert_inline_picture_with_docx(
                    word_abs_path=word_abs_path,
                    image_abs_path=image_abs_path,
                    width=width,
                    height=height,
                    target_text=target_text
                )
            
            # 非嵌入型图片依赖 Word COM，只能在 Windows + Word 环境下工作。
            if not WIN32COM_AVAILABLE:
                return {
                    "success": False,
                    "error": "当前环境不支持非嵌入型图片的文字环绕/浮动定位。请在 Windows + Word + pywin32 环境下使用，或改用嵌入型。"
                }
            
            word = None
            doc = None
            
            # 创建Word应用程序实例
            word = win32com.client.Dispatch("Word.Application")
            word.Visible = False
            word.DisplayAlerts = False
            
            doc = word.Documents.Open(str(word_abs_path))
            
            if target_text:
                search_range = doc.Range()
                search_range.Find.Execute(FindText=target_text)
                
                if not search_range.Find.Found:
                    doc.Close(SaveChanges=False)
                    word.Quit()
                    return {"success": False, "error": f"未找到目标文本: {target_text}"}
                
                search_range.Select()
                selection = word.Selection
            else:
                selection = word.Selection
                selection.Collapse(1)
            
            inline_shape = selection.InlineShapes.AddPicture(
                FileName=str(image_abs_path),
                LinkToFile=False,
                SaveWithDocument=True
            )
            
            shape = inline_shape.ConvertToShape()
            shape.WrapFormat.Type = wrap_value
            
            if wrap_value == 5:
                shape.ZOrder(5)
            elif wrap_value == 4:
                shape.ZOrder(0)
            
            shape.WrapFormat.AllowOverlap = True
            
            if position_x is not None:
                shape.Left = position_x
            if position_y is not None:
                shape.Top = position_y
            
            if width is not None or height is not None:
                shape.LockAspectRatio = True
                if width is not None:
                    shape.Width = width
                if height is not None:
                    shape.Height = height
            
            doc.Save()
            doc.Close(SaveChanges=False)
            word.Quit()
            
            # 验证文件大小是否增加
            file_size_after = word_abs_path.stat().st_size
            if file_size_after <= file_size_before:
                return {
                    "success": False,
                    "error": f"图片插入失败：文件大小未增加（插入前: {file_size_before} 字节，插入后: {file_size_after} 字节）"
                }
            
            wrap_type_name = next((k for k, v in wrap_type_map.items() if v == wrap_value and len(k) > 2), str(wrap_value))
            
            return {
                "success": True,
                "message": f"成功插入图片并设置为{wrap_type_name}",
                "wrap_type": wrap_type_name,
                "output_file": str(word_abs_path),
                "file_size_before": file_size_before,
                "file_size_after": file_size_after,
                "size_increase": file_size_after - file_size_before,
                "backend": "win32com"
            }
                
        except Exception as e:
            # 确保清理资源
            try:
                if doc is not None:
                    doc.Close(SaveChanges=False)
            except:
                pass
            
            try:
                if word is not None:
                    word.Quit()
            except:
                pass
            
            return {"success": False, "error": f"插入图片失败: {str(e)}"}

    # ════════════════════════════════════════════════════════════
    # 一、论文模板与页面样式初始化
    # ════════════════════════════════════════════════════════════

    def initialize_academic_template(
        self,
        title: str = "",
        author: str = "",
        institution: str = "",
        date_text: str = "",
        language: str = "zh",
        paper_type: str = "article",
        style_preset: str = "cn_academic",
        include_cover: bool = True,
        include_abstract_placeholders: bool = True,
        include_toc_placeholder: bool = True,
    ) -> Dict[str, Any]:
        """
        一键创建适合论文写作的文档基线。
        设置页面大小、页边距、默认样式、标题层级样式，
        并可选生成封面、摘要占位、目录占位、章节占位。

        :param title: 论文标题
        :param author: 作者
        :param institution: 机构
        :param date_text: 日期文本
        :param language: "zh" / "en"
        :param paper_type: "article" / "thesis" / "report"
        :param style_preset: "cn_academic" / "ieee_like" / "apa_like"
        :param include_cover: 是否生成封面
        :param include_abstract_placeholders: 是否生成摘要占位章节
        :param include_toc_placeholder: 是否插入目录占位
        :return: 初始化结果及已创建章节列表
        """
        if not DOCX_AVAILABLE:
            return {"success": False, "error": "python-docx 库不可用，无法初始化学术模板"}

        try:
            self.document = Document()
            preset = self.ACADEMIC_PRESETS.get(style_preset, self.ACADEMIC_PRESETS["cn_academic"])
            created_sections = []

            # ── 1. 页面设置 ──
            section = self.document.sections[0]
            section.page_width = Cm(21.0)
            section.page_height = Cm(29.7)
            margins = preset["margins_cm"]
            section.top_margin = Cm(margins["top"])
            section.bottom_margin = Cm(margins["bottom"])
            section.left_margin = Cm(margins["left"])
            section.right_margin = Cm(margins["right"])

            # ── 2. 正文默认样式 ──
            style_normal = self.document.styles['Normal']
            font = style_normal.font
            font.name = preset["body_font"]
            font.size = Pt(preset["body_font_size"])
            rpr = style_normal.element.get_or_add_rPr()
            rFonts = rpr.find(qn('w:rFonts'))
            if rFonts is None:
                rFonts = OxmlElement('w:rFonts')
                rpr.append(rFonts)
            rFonts.set(qn('w:eastAsia'), preset["body_font"])
            pf = style_normal.paragraph_format
            pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
            pf.line_spacing = preset["body_line_spacing"]
            if preset["body_first_line_indent_cm"] > 0:
                pf.first_line_indent = Cm(preset["body_first_line_indent_cm"])

            # ── 3. 标题层级样式 ──
            for level, hs in preset.get("heading_styles", {}).items():
                style_name = f"Heading {level}"
                if style_name in self.document.styles:
                    h_style = self.document.styles[style_name]
                else:
                    h_style = self.document.styles.add_style(style_name, WD_STYLE_TYPE.PARAGRAPH)
                h_font = h_style.font
                h_font.name = hs["font"]
                h_font.size = Pt(hs["size"])
                h_font.bold = hs.get("bold", True)
                if hs.get("italic"):
                    h_font.italic = True
                h_rpr = h_style.element.get_or_add_rPr()
                h_rFonts = h_rpr.find(qn('w:rFonts'))
                if h_rFonts is None:
                    h_rFonts = OxmlElement('w:rFonts')
                    h_rpr.append(h_rFonts)
                h_rFonts.set(qn('w:eastAsia'), hs["font"])
                align = hs.get("alignment", "left")
                if align == "center":
                    h_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
                elif align == "right":
                    h_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                else:
                    h_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
                h_style.paragraph_format.space_before = Pt(12)
                h_style.paragraph_format.space_after = Pt(6)

            # ── 4. 封面 ──
            if include_cover:
                if title:
                    p_title = self.document.add_paragraph()
                    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run_t = p_title.add_run(title)
                    run_t.bold = True
                    run_t.font.size = Pt(22)
                    run_t.font.name = preset["heading_styles"].get(1, {}).get("font", preset["body_font"])
                if author:
                    p_author = self.document.add_paragraph()
                    p_author.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run_a = p_author.add_run(author)
                    run_a.font.size = Pt(14)
                if institution:
                    p_inst = self.document.add_paragraph()
                    p_inst.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run_i = p_inst.add_run(institution)
                    run_i.font.size = Pt(12)
                if date_text:
                    p_date = self.document.add_paragraph()
                    p_date.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run_d = p_date.add_run(date_text)
                    run_d.font.size = Pt(12)
                self.document.add_page_break()

            # ── 5. 目录占位 ──
            if include_toc_placeholder:
                toc_label = "目录" if language == "zh" else "Table of Contents"
                self.document.add_heading(toc_label, level=1)
                p_toc = self.document.add_paragraph()
                run_placeholder = p_toc.add_run('【目录占位区 — 请在 Word 中打开后右键"更新域"以生成实际目录】')
                run_placeholder.font.color.rgb = RGBColor(128, 128, 128)
                run_placeholder.font.italic = True
                self._insert_toc_field_xml(p_toc)
                self.document.add_page_break()
                created_sections.append(toc_label)

            # ── 6. 占位章节 ──
            if include_abstract_placeholders:
                if paper_type == "thesis":
                    sections_template = self.THESIS_SECTIONS.get(language, self.THESIS_SECTIONS["zh"])
                else:
                    sections_template = self.ARTICLE_SECTIONS.get(language, self.ARTICLE_SECTIONS["zh"])
                for sec_title, level in sections_template:
                    if level > 0:
                        self.document.add_heading(sec_title, level=level)
                        created_sections.append(sec_title)
                    else:
                        p_kw = self.document.add_paragraph()
                        run_kw_label = p_kw.add_run(sec_title + "：")
                        run_kw_label.bold = True
                        run_kw_content = p_kw.add_run("【待填写】")
                        run_kw_content.font.color.rgb = RGBColor(128, 128, 128)

            # 重置学术计数器
            self._caption_counters = {}
            self._reference_entries = {}
            self._reference_order = []
            self._review_counter = 0
            self._todo_counter = 0

            return {
                "success": True,
                "message": f"学术模板初始化完成（预设: {style_preset}，语言: {language}，类型: {paper_type}）",
                "preset": style_preset,
                "language": language,
                "paper_type": paper_type,
                "created_sections": created_sections,
                "note": "目录需在 Word 中手动更新域" if include_toc_placeholder else "",
            }
        except Exception as e:
            return {"success": False, "error": f"初始化学术模板失败: {str(e)}"}

    def _insert_toc_field_xml(self, placeholder_paragraph):
        """
        尝试在目录占位段落中插入 TOC 域代码（尽力而为，不影响跨平台基本功能）。
        """
        try:
            fld_char_begin = OxmlElement('w:fldChar')
            fld_char_begin.set(qn('w:fldCharType'), 'begin')
            fld_code = OxmlElement('w:instrText')
            fld_code.set(qn('xml:space'), 'preserve')
            fld_code.text = ' TOC \\o "1-3" \\h \\z \\u '
            fld_char_separate = OxmlElement('w:fldChar')
            fld_char_separate.set(qn('w:fldCharType'), 'separate')
            fld_char_end = OxmlElement('w:fldChar')
            fld_char_end.set(qn('w:fldCharType'), 'end')

            run1 = placeholder_paragraph.add_run()._element
            run2 = placeholder_paragraph.add_run()._element
            run3 = placeholder_paragraph.add_run()._element
            run4 = placeholder_paragraph.add_run()._element
            run1.append(fld_char_begin)
            run2.append(fld_code)
            run3.append(fld_char_separate)
            run4.append(fld_char_end)
        except Exception:
            pass

    # ════════════════════════════════════════════════════════════
    # 二、题注、自动编号、目录/图表目录
    # ════════════════════════════════════════════════════════════

    def insert_caption(
        self,
        target_type: str = "figure",
        caption_text: str = "",
        numbering_scope: str = "global",
        position: str = "after",
        target_index: Optional[int] = None,
    ) -> Dict[str, Any]:
        """
        插入题注（图/表/公式），自动生成编号如"图 1 xxx""表 2 xxx"。

        :param target_type: "figure" / "table" / "equation"
        :param caption_text: 题注文字
        :param numbering_scope: 编号范围（目前仅支持 "global" 全文档递增）
        :param position: "before" / "after"
        :param target_index: 可选，目标段落索引
        :return: 含编号和题注文本的结果
        """
        if not DOCX_AVAILABLE:
            return {"success": False, "error": "python-docx 库不可用，无法插入题注"}

        try:
            if not self.document:
                return {"success": False, "error": "文档未加载"}

            preset_name = getattr(self, '_active_preset', 'cn_academic')
            preset = self.ACADEMIC_PRESETS.get(preset_name, self.ACADEMIC_PRESETS["cn_academic"])
            label_map = {
                "figure": preset.get("label_figure", "图"),
                "table": preset.get("label_table", "表"),
                "equation": preset.get("label_equation", "式"),
            }
            if target_type not in label_map:
                return {"success": False, "error": "target_type 必须是 figure、table 或 equation"}
            if numbering_scope != "global":
                return {"success": False, "error": "当前仅支持 global 编号范围"}
            if position not in ("before", "after"):
                return {"success": False, "error": "position 必须是 before 或 after"}
            if target_index is not None and (target_index < 0 or target_index >= len(self.document.paragraphs)):
                return {"success": False, "error": f"目标段落索引 {target_index} 超出范围"}

            label = label_map[target_type]

            self._sync_caption_counters_from_document()
            self._caption_counters.setdefault(target_type, 0)
            self._caption_counters[target_type] += 1
            number = self._caption_counters[target_type]

            full_caption = f"{label} {number} {caption_text}".strip()

            p = self.document.add_paragraph()
            try:
                p.style = "Caption"
            except Exception:
                pass
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.add_run(f"{label} ")
            self._add_seq_field(p, self._caption_sequence_name(target_type), str(number))
            if caption_text:
                p.add_run(f" {caption_text}")
            for run in p.runs:
                run.font.size = Pt(10)
                run.font.name = preset.get("body_font", "宋体")

            if target_index is not None:
                body = self.document.element.body
                body.remove(p._element)
                ref_para = self.document.paragraphs[target_index]
                if position == "before":
                    ref_para._element.addprevious(p._element)
                else:
                    ref_para._element.addnext(p._element)

            return {
                "success": True,
                "caption": full_caption,
                "label": label,
                "number": number,
                "target_type": target_type,
                "message": f"题注插入成功: {full_caption}",
            }
        except Exception as e:
            return {"success": False, "error": f"插入题注失败: {str(e)}"}

    def insert_table_with_caption(
        self,
        rows: int,
        cols: int,
        caption_text: str = "",
        headers: Optional[List[str]] = None,
        data: Optional[List[List[str]]] = None,
        position: int = -1,
    ) -> Dict[str, Any]:
        """
        封装"建表 + 题注"：先插入表格，再在表格上方插入题注。

        :param rows: 行数
        :param cols: 列数
        :param caption_text: 题注文字
        :param headers: 表头列表
        :param data: 数据二维列表
        :param position: 插入位置，-1 为末尾
        """
        if not DOCX_AVAILABLE:
            return {"success": False, "error": "python-docx 库不可用"}

        try:
            if not self.document:
                return {"success": False, "error": "文档未加载"}
            if position < -1:
                return {"success": False, "error": f"插入位置 {position} 无效，必须大于等于 -1"}

            caption_target_index = position if 0 <= position < len(self.document.paragraphs) else None
            cap_result = self.insert_caption(
                target_type="table",
                caption_text=caption_text,
                position="before",
                target_index=caption_target_index,
            )
            if not cap_result["success"]:
                return cap_result

            tbl_result = self.create_table(rows, cols, headers=headers, data=data, position=position)
            if not tbl_result["success"]:
                return tbl_result

            return {
                "success": True,
                "message": f"带题注表格插入成功: {cap_result['caption']}",
                "caption": cap_result["caption"],
                "table_rows": rows,
                "table_cols": cols,
            }
        except Exception as e:
            return {"success": False, "error": f"插入带题注表格失败: {str(e)}"}

    def insert_picture_with_caption(
        self,
        word_path: str,
        image_path: str,
        caption_text: str = "",
        wrap_type: str = "嵌入型",
        width: Optional[float] = None,
        height: Optional[float] = None,
        target_text: Optional[str] = None,
    ) -> Dict[str, Any]:
        """
        封装"插图 + 题注"：先插入图片，再在图片下方插入题注。

        :param word_path: Word 文档路径
        :param image_path: 图片路径
        :param caption_text: 题注文字
        :param wrap_type: 文字环绕类型
        :param width: 宽度
        :param height: 高度
        :param target_text: 目标文本
        """
        try:
            word_abs_path = self._resolve_path(word_path)
            if self.document:
                save_result = self.save_document(str(word_abs_path))
                if not save_result["success"]:
                    return save_result

            pic_result = self.insert_picture_with_wrap(
                word_path=word_path,
                image_path=image_path,
                wrap_type=wrap_type,
                width=width,
                height=height,
                target_text=target_text,
            )
            if not pic_result["success"]:
                return pic_result

            reload_result = self.load_document(str(word_abs_path))
            if not reload_result["success"]:
                return {
                    "success": True,
                    "message": f"图片插入成功但重新加载文档失败: {reload_result.get('error', '')}",
                    "reload_error": reload_result.get("error"),
                    "picture_result": pic_result,
                }

            inserted_paragraph_index = pic_result.get("inserted_paragraph_index")
            cap_result = self.insert_caption(
                target_type="figure",
                caption_text=caption_text,
                position="after",
                target_index=inserted_paragraph_index if isinstance(inserted_paragraph_index, int) else None,
            )
            if not cap_result["success"]:
                return {
                    "success": True,
                    "message": f"图片插入成功但题注失败: {cap_result.get('error', '')}",
                    "caption_error": cap_result.get("error"),
                    "picture_result": pic_result,
                }

            save_result = self.save_document(str(word_abs_path))
            if not save_result["success"]:
                return {
                    "success": True,
                    "message": f"图片和题注已写入内存但保存失败: {save_result.get('error', '')}",
                    "caption": cap_result["caption"],
                    "save_error": save_result.get("error"),
                    "picture_result": pic_result,
                }

            return {
                "success": True,
                "message": f"带题注图片插入成功: {cap_result['caption']}",
                "caption": cap_result["caption"],
                "picture_result": pic_result,
                "save_result": save_result,
            }
        except Exception as e:
            return {"success": False, "error": f"插入带题注图片失败: {str(e)}"}

    def insert_toc_placeholder(self, language: str = "zh") -> Dict[str, Any]:
        """
        在当前位置插入目录占位区。
        跨平台下插入占位文字；同时尝试插入 TOC 域代码。

        :param language: "zh" / "en"
        """
        if not DOCX_AVAILABLE:
            return {"success": False, "error": "python-docx 库不可用"}

        try:
            if not self.document:
                return {"success": False, "error": "文档未加载"}

            label = "目录" if language == "zh" else "Table of Contents"
            self.document.add_heading(label, level=1)
            p_toc = self.document.add_paragraph()
            run_ph = p_toc.add_run('【目录占位区 — 请在 Word 中打开后右键"更新域"以生成实际目录】')
            run_ph.font.color.rgb = RGBColor(128, 128, 128)
            run_ph.font.italic = True
            self._insert_toc_field_xml(p_toc)
            self.document.add_page_break()
            return {
                "success": True,
                "message": "目录占位区已插入",
                "note": "需在 Word 中打开后更新域以生成实际目录",
            }
        except Exception as e:
            return {"success": False, "error": f"插入目录占位失败: {str(e)}"}

    def insert_list_of_figures_placeholder(self, language: str = "zh") -> Dict[str, Any]:
        """
        插入图表目录（图目录）占位区。

        :param language: "zh" / "en"
        """
        if not DOCX_AVAILABLE:
            return {"success": False, "error": "python-docx 库不可用"}

        try:
            if not self.document:
                return {"success": False, "error": "文档未加载"}

            label = "图目录" if language == "zh" else "List of Figures"
            self.document.add_heading(label, level=1)
            p = self.document.add_paragraph()
            run_ph = p.add_run("【图目录占位区 — 需在 Word 中更新域】")
            run_ph.font.color.rgb = RGBColor(128, 128, 128)
            run_ph.font.italic = True
            self.document.add_page_break()
            return {
                "success": True,
                "message": "图目录占位区已插入",
                "note": "需在 Word 中打开后更新域",
            }
        except Exception as e:
            return {"success": False, "error": f"插入图目录占位失败: {str(e)}"}

    def insert_list_of_tables_placeholder(self, language: str = "zh") -> Dict[str, Any]:
        """
        插入表目录占位区。

        :param language: "zh" / "en"
        """
        if not DOCX_AVAILABLE:
            return {"success": False, "error": "python-docx 库不可用"}

        try:
            if not self.document:
                return {"success": False, "error": "文档未加载"}

            label = "表目录" if language == "zh" else "List of Tables"
            self.document.add_heading(label, level=1)
            p = self.document.add_paragraph()
            run_ph = p.add_run("【表目录占位区 — 需在 Word 中更新域】")
            run_ph.font.color.rgb = RGBColor(128, 128, 128)
            run_ph.font.italic = True
            self.document.add_page_break()
            return {
                "success": True,
                "message": "表目录占位区已插入",
                "note": "需在 Word 中打开后更新域",
            }
        except Exception as e:
            return {"success": False, "error": f"插入表目录占位失败: {str(e)}"}

    # ════════════════════════════════════════════════════════════
    # 三、参考文献与文内引用管理
    # ════════════════════════════════════════════════════════════

    def add_reference_entry(
        self,
        cite_key: str,
        authors: str = "",
        title: str = "",
        year: str = "",
        journal: str = "",
        volume: str = "",
        issue: str = "",
        pages: str = "",
        publisher: str = "",
        doi: str = "",
        url: str = "",
        entry_type: str = "article",
        style: str = "gbt7714",
    ) -> Dict[str, Any]:
        """
        登记一条参考文献条目到文档级缓存。

        :param cite_key: 引用键（唯一标识）
        :param authors: 作者
        :param title: 标题
        :param year: 年份
        :param journal: 期刊/书名
        :param volume: 卷
        :param issue: 期
        :param pages: 页码
        :param publisher: 出版社
        :param doi: DOI
        :param url: URL
        :param entry_type: article / book / conference / webpage / thesis / misc
        :param style: gbt7714 / ieee / apa_like
        :return: 登记结果
        """
        try:
            if cite_key in self._reference_entries:
                return {
                    "success": False,
                    "error": f"引用键 '{cite_key}' 已存在，请使用不同的键",
                }

            entry = {
                "cite_key": cite_key,
                "authors": authors,
                "title": title,
                "year": year,
                "journal": journal,
                "volume": volume,
                "issue": issue,
                "pages": pages,
                "publisher": publisher,
                "doi": doi,
                "url": url,
                "entry_type": entry_type,
                "style": style,
            }
            self._reference_entries[cite_key] = entry
            return {
                "success": True,
                "message": f"参考文献条目 '{cite_key}' 登记成功",
                "cite_key": cite_key,
                "entry_type": entry_type,
            }
        except Exception as e:
            return {"success": False, "error": f"登记参考文献失败: {str(e)}"}

    def _format_reference(self, entry: Dict[str, Any], number: int) -> str:
        """根据条目类型和样式格式化单条参考文献。"""
        style = entry.get("style", "gbt7714")
        etype = entry.get("entry_type", "article")
        authors = entry.get("authors", "")
        title = entry.get("title", "")
        year = entry.get("year", "")
        journal = entry.get("journal", "")
        volume = entry.get("volume", "")
        issue = entry.get("issue", "")
        pages = entry.get("pages", "")
        publisher = entry.get("publisher", "")
        doi = entry.get("doi", "")
        url = entry.get("url", "")

        if style == "ieee":
            parts = [f"[{number}] "]
            if authors:
                parts.append(f"{authors}, ")
            parts.append(f"\"{title},\" ")
            if journal:
                parts.append(f"{journal}, ")
            if volume:
                parts.append(f"vol. {volume}, ")
            if issue:
                parts.append(f"no. {issue}, ")
            if pages:
                parts.append(f"pp. {pages}, ")
            if year:
                parts.append(f"{year}.")
            if doi:
                parts.append(f" doi: {doi}")
            return "".join(parts)

        elif style == "apa_like":
            parts = []
            if authors:
                parts.append(f"{authors} ")
            if year:
                parts.append(f"({year}). ")
            parts.append(f"{title}. ")
            if journal:
                parts.append(f"{journal}, ")
            if volume:
                parts.append(f"{volume}")
            if issue:
                parts.append(f"({issue})")
            if pages:
                parts.append(f", {pages}.")
            else:
                parts.append(".")
            if doi:
                parts.append(f" https://doi.org/{doi}")
            return "".join(parts)

        else:
            # GB/T 7714 风格（简化）
            type_markers = {
                "article": "[J]",
                "book": "[M]",
                "conference": "[C]",
                "webpage": "[EB/OL]",
                "thesis": "[D]",
                "misc": "[Z]",
            }
            marker = type_markers.get(etype, "[J]")
            parts = [f"[{number}] "]
            if authors:
                parts.append(f"{authors}. ")
            parts.append(f"{title}{marker}. ")
            if journal:
                parts.append(f"{journal}, ")
            if volume:
                parts.append(f"{volume}")
            if issue:
                parts.append(f"({issue})")
            if pages:
                parts.append(f": {pages}")
            if year:
                parts.append(f", {year}.")
            if url:
                parts.append(f" {url}")
            return "".join(parts)

    def insert_citation(
        self,
        cite_keys: List[str],
        style: str = "gbt7714",
        position_mode: str = "append_current_paragraph",
        target_text: Optional[str] = None,
    ) -> Dict[str, Any]:
        """
        在正文中插入轻量引用标记。
        支持 IEEE 风格 [1][2]、author-year 风格 (Author, 2024)、简化 GB/T [1]。

        :param cite_keys: 引用键列表
        :param style: "gbt7714" / "ieee" / "apa_like"
        :param position_mode: "append_current_paragraph" / "insert_new_paragraph" / "target_text"
        :param target_text: 当 position_mode 为 "target_text" 时使用
        :return: 插入结果
        """
        if not DOCX_AVAILABLE:
            return {"success": False, "error": "python-docx 库不可用"}

        try:
            if not self.document:
                return {"success": False, "error": "文档未加载"}

            if not cite_keys:
                return {"success": False, "error": "cite_keys 不能为空"}

            numbers = []
            for key in cite_keys:
                if key not in self._reference_entries:
                    return {"success": False, "error": f"引用键 '{key}' 未登记，请先调用 add_reference_entry"}
                if key not in self._reference_order:
                    self._reference_order.append(key)
                numbers.append(self._reference_order.index(key) + 1)

            if style in ("gbt7714", "ieee"):
                citation_text = "".join(f"[{n}]" for n in numbers)
            elif style == "apa_like":
                parts = []
                for key in cite_keys:
                    entry = self._reference_entries[key]
                    author_short = entry.get("authors", "").split(",")[0].split(" ")[0] if entry.get("authors") else "Unknown"
                    yr = entry.get("year", "n.d.")
                    parts.append(f"{author_short}, {yr}")
                citation_text = "(" + "; ".join(parts) + ")"
            else:
                citation_text = "".join(f"[{n}]" for n in numbers)

            if position_mode == "target_text" and target_text:
                found = False
                for para in self.document.paragraphs:
                    if target_text in para.text:
                        run = para.add_run(citation_text)
                        run.font.size = Pt(10)
                        found = True
                        break
                if not found:
                    return {"success": False, "error": f"未找到目标文本: {target_text}"}
            elif position_mode == "insert_new_paragraph":
                p = self.document.add_paragraph()
                run = p.add_run(citation_text)
                run.font.size = Pt(10)
            else:
                if self.document.paragraphs:
                    last_para = self.document.paragraphs[-1]
                    run = last_para.add_run(citation_text)
                    run.font.size = Pt(10)
                else:
                    p = self.document.add_paragraph()
                    run = p.add_run(citation_text)
                    run.font.size = Pt(10)

            return {
                "success": True,
                "citation_text": citation_text,
                "cite_keys": cite_keys,
                "numbers": numbers,
                "style": style,
                "message": f"引用标记 '{citation_text}' 插入成功",
            }
        except Exception as e:
            return {"success": False, "error": f"插入引用失败: {str(e)}"}

    def generate_bibliography(
        self,
        style: str = "gbt7714",
        heading_text: str = "",
        sort_mode: str = "appearance",
    ) -> Dict[str, Any]:
        """
        根据已登记的 reference entries 自动生成"参考文献"章节。

        :param style: "gbt7714" / "ieee" / "apa_like"
        :param heading_text: 参考文献标题
        :param sort_mode: "appearance" / "alphabetical"
        :return: 生成结果
        """
        if not DOCX_AVAILABLE:
            return {"success": False, "error": "python-docx 库不可用"}

        try:
            if not self.document:
                return {"success": False, "error": "文档未加载"}

            if not self._reference_entries:
                return {"success": False, "error": "尚未登记任何参考文献条目"}

            if sort_mode == "alphabetical":
                sorted_keys = sorted(self._reference_entries.keys())
            else:
                sorted_keys = []
                for key in self._reference_order:
                    if key in self._reference_entries:
                        sorted_keys.append(key)
                for key in self._reference_entries:
                    if key not in sorted_keys:
                        sorted_keys.append(key)

            if not heading_text:
                heading_text = "参考文献"
            self.document.add_heading(heading_text, level=1)

            entries_info = []
            for i, key in enumerate(sorted_keys, 1):
                entry = self._reference_entries[key]
                formatted = self._format_reference(entry, i)
                p = self.document.add_paragraph()
                run = p.add_run(formatted)
                run.font.size = Pt(10.5)
                entries_info.append({"cite_key": key, "number": i, "formatted": formatted})

            return {
                "success": True,
                "message": f"参考文献列表生成成功，共 {len(entries_info)} 条",
                "count": len(entries_info),
                "entries": entries_info,
            }
        except Exception as e:
            return {"success": False, "error": f"生成参考文献列表失败: {str(e)}"}

    # ════════════════════════════════════════════════════════════
    # 四、章节插入与结构重组
    # ════════════════════════════════════════════════════════════

    def insert_heading_after_text(
        self,
        text_anchor: str,
        heading_text: str,
        level: int = 1,
    ) -> Dict[str, Any]:
        """
        在指定文本锚点后插入标题。

        :param text_anchor: 锚点文本
        :param heading_text: 标题文本
        :param level: 标题级别 1-9
        :return: 插入结果
        """
        if not DOCX_AVAILABLE:
            return {"success": False, "error": "python-docx 库不可用"}

        try:
            if not self.document:
                return {"success": False, "error": "文档未加载"}

            if level < 1 or level > 9:
                return {"success": False, "error": "标题级别必须在1-9之间"}

            anchor_para = None
            for para in self.document.paragraphs:
                if text_anchor in para.text:
                    anchor_para = para
                    break

            if anchor_para is None:
                return {"success": False, "error": f"未找到锚点文本: {text_anchor}"}

            new_para = OxmlElement('w:p')
            anchor_para._element.addnext(new_para)
            temp_para = docx.text.paragraph.Paragraph(new_para, self.document)
            temp_para.style = self.document.styles[f"Heading {level}"]
            temp_para.add_run(heading_text)

            return {
                "success": True,
                "message": f"在 '{text_anchor}' 后插入 {level} 级标题 '{heading_text}' 成功",
                "heading_text": heading_text,
                "level": level,
            }
        except Exception as e:
            return {"success": False, "error": f"插入标题失败: {str(e)}"}

    def insert_section_block(
        self,
        heading_text: str,
        level: int = 1,
        paragraphs: Optional[List[str]] = None,
        after_heading: Optional[str] = None,
        after_text: Optional[str] = None,
        page_break_before: bool = False,
    ) -> Dict[str, Any]:
        """
        一次插入一整节：标题 + 多个段落。

        :param heading_text: 标题文本
        :param level: 标题级别
        :param paragraphs: 段落内容列表
        :param after_heading: 在此标题文本之后插入
        :param after_text: 在此正文文本之后插入（优先级低于 after_heading）
        :param page_break_before: 是否在节前插入分页符
        :return: 插入结果
        """
        if not DOCX_AVAILABLE:
            return {"success": False, "error": "python-docx 库不可用"}

        try:
            if not self.document:
                return {"success": False, "error": "文档未加载"}

            if level < 1 or level > 9:
                return {"success": False, "error": "标题级别必须在1-9之间"}

            anchor_para = None
            if after_heading:
                for para in self.document.paragraphs:
                    if para.style.name.startswith("Heading") and after_heading in para.text:
                        anchor_para = para
                        break
            if anchor_para is None and after_text:
                for para in self.document.paragraphs:
                    if after_text in para.text:
                        anchor_para = para
                        break

            insert_at_end = anchor_para is None

            if page_break_before:
                if insert_at_end:
                    self.document.add_page_break()
                else:
                    pb_para = anchor_para.insert_paragraph_before("")
                    run_pb = pb_para.add_run()
                    run_pb.add_break(docx.enum.text.WD_BREAK.PAGE)
                    anchor_para._element.addnext(pb_para._element)

            if insert_at_end:
                self.document.add_heading(heading_text, level=level)
            else:
                new_heading_para = OxmlElement('w:p')
                anchor_para._element.addnext(new_heading_para)
                temp_para = docx.text.paragraph.Paragraph(new_heading_para, self.document)
                temp_para.style = self.document.styles[f"Heading {level}"]
                temp_para.add_run(heading_text)
                anchor_para = temp_para

            inserted_count = 0
            if paragraphs:
                for para_text in paragraphs:
                    if insert_at_end:
                        self.document.add_paragraph(para_text)
                    else:
                        new_p = OxmlElement('w:p')
                        anchor_para._element.addnext(new_p)
                        temp_p = docx.text.paragraph.Paragraph(new_p, self.document)
                        temp_p.add_run(para_text)
                        anchor_para = temp_p
                    inserted_count += 1

            return {
                "success": True,
                "message": f"节块 '{heading_text}' 插入成功（含 {inserted_count} 个段落）",
                "heading_text": heading_text,
                "level": level,
                "paragraphs_inserted": inserted_count,
            }
        except Exception as e:
            return {"success": False, "error": f"插入节块失败: {str(e)}"}

    def move_section(
        self,
        section_heading: str,
        target_after_heading: str,
    ) -> Dict[str, Any]:
        """
        轻量实现章节块移动：复制到新位置 + 删除旧块。
        块范围：从 section_heading 到下一个同级或更高级 Heading 之间。

        :param section_heading: 要移动的章节标题文本
        :param target_after_heading: 目标位置标题
        :return: 移动结果
        """
        if not DOCX_AVAILABLE:
            return {"success": False, "error": "python-docx 库不可用"}

        try:
            if not self.document:
                return {"success": False, "error": "文档未加载"}

            source_start_idx = None
            source_level = None
            for i, para in enumerate(self.document.paragraphs):
                if para.style.name.startswith("Heading") and section_heading in para.text:
                    source_start_idx = i
                    try:
                        source_level = int(para.style.name.split()[-1])
                    except (ValueError, IndexError):
                        source_level = 1
                    break

            if source_start_idx is None:
                return {"success": False, "error": f"未找到源章节标题: {section_heading}"}

            source_end_idx = len(self.document.paragraphs)
            for i in range(source_start_idx + 1, len(self.document.paragraphs)):
                para = self.document.paragraphs[i]
                if para.style.name.startswith("Heading"):
                    try:
                        plevel = int(para.style.name.split()[-1])
                    except (ValueError, IndexError):
                        plevel = 9
                    if plevel <= source_level:
                        source_end_idx = i
                        break

            block_content = []
            for i in range(source_start_idx, source_end_idx):
                para = self.document.paragraphs[i]
                style_name = para.style.name if para.style else "Normal"
                text = para.text
                block_content.append({"style": style_name, "text": text})

            target_idx = None
            for i, para in enumerate(self.document.paragraphs):
                if para.style.name.startswith("Heading") and target_after_heading in para.text:
                    try:
                        tgt_level = int(para.style.name.split()[-1])
                    except (ValueError, IndexError):
                        tgt_level = 1
                    target_idx = i
                    for j in range(i + 1, len(self.document.paragraphs)):
                        p = self.document.paragraphs[j]
                        if p.style.name.startswith("Heading"):
                            try:
                                pl = int(p.style.name.split()[-1])
                            except (ValueError, IndexError):
                                pl = 9
                            if pl <= tgt_level:
                                target_idx = j
                                break
                        target_idx = j + 1
                    break

            if target_idx is None:
                return {"success": False, "error": f"未找到目标章节标题: {target_after_heading}"}

            target_para = self.document.paragraphs[min(target_idx, len(self.document.paragraphs) - 1)]
            for item in reversed(block_content):
                new_p = OxmlElement('w:p')
                target_para._element.addnext(new_p)
                temp_p = docx.text.paragraph.Paragraph(new_p, self.document)
                if item["style"] in [s.name for s in self.document.styles]:
                    temp_p.style = self.document.styles[item["style"]]
                temp_p.add_run(item["text"])

            # 删除旧块
            to_delete = []
            for i, para in enumerate(self.document.paragraphs):
                if para.style.name.startswith("Heading") and section_heading in para.text:
                    src_lvl = None
                    try:
                        src_lvl = int(para.style.name.split()[-1])
                    except (ValueError, IndexError):
                        src_lvl = 1
                    for j in range(i, len(self.document.paragraphs)):
                        p = self.document.paragraphs[j]
                        if j > i and p.style.name.startswith("Heading"):
                            try:
                                pl = int(p.style.name.split()[-1])
                            except (ValueError, IndexError):
                                pl = 9
                            if pl <= src_lvl:
                                break
                        to_delete.append(p)
                    break

            for p in to_delete:
                p._element.getparent().remove(p._element)

            return {
                "success": True,
                "message": f"章节 '{section_heading}' 已移动到 '{target_after_heading}' 之后",
                "block_size": len(block_content),
            }
        except Exception as e:
            return {"success": False, "error": f"移动章节失败: {str(e)}"}

    def get_document_outline(self) -> Dict[str, Any]:
        """
        返回按标题层级组织的文档大纲。

        :return: 文档大纲
        """
        if not DOCX_AVAILABLE:
            return {"success": False, "error": "python-docx 库不可用"}

        try:
            if not self.document:
                return {"success": False, "error": "文档未加载"}

            outline = []
            for i, para in enumerate(self.document.paragraphs):
                if para.style and para.style.name.startswith("Heading"):
                    try:
                        level = int(para.style.name.split()[-1])
                    except (ValueError, IndexError):
                        level = 1
                    outline.append({
                        "index": i,
                        "level": level,
                        "text": para.text,
                    })

            return {
                "success": True,
                "outline": outline,
                "heading_count": len(outline),
            }
        except Exception as e:
            return {"success": False, "error": f"获取文档大纲失败: {str(e)}"}

    # ════════════════════════════════════════════════════════════
    # 五、审阅/修订辅助能力
    # ════════════════════════════════════════════════════════════

    def insert_review_comment_block(
        self,
        comment_text: str,
        comment_type: str = "todo",
        target_text: Optional[str] = None,
        after_paragraph_index: Optional[int] = None,
    ) -> Dict[str, Any]:
        """
        在文中插入格式化审阅块（跨平台稳定版，非 Word 原生批注）。

        :param comment_text: 审阅内容
        :param comment_type: todo / citation_needed / rewrite / logic_check / format_check
        :param target_text: 锚点文本
        :param after_paragraph_index: 在第 N 个段落后插入
        :return: 插入结果
        """
        if not DOCX_AVAILABLE:
            return {"success": False, "error": "python-docx 库不可用"}

        try:
            if not self.document:
                return {"success": False, "error": "文档未加载"}

            self._review_counter += 1
            review_id = self._review_counter

            type_labels = {
                "todo": "TODO",
                "citation_needed": "CITATION_NEEDED",
                "rewrite": "REWRITE",
                "logic_check": "LOGIC_CHECK",
                "format_check": "FORMAT_CHECK",
            }
            label = type_labels.get(comment_type, "REVIEW")
            marker_text = f"[REVIEW][{label}] {comment_text}"

            p = self.document.add_paragraph()
            run = p.add_run(marker_text)
            run.bold = True
            run.font.color.rgb = RGBColor(255, 0, 0)
            run.font.size = Pt(10)

            if target_text:
                for para in self.document.paragraphs:
                    if target_text in para.text and para != p:
                        para._element.addnext(p._element)
                        break
            elif after_paragraph_index is not None and 0 <= after_paragraph_index < len(self.document.paragraphs):
                ref_para = self.document.paragraphs[after_paragraph_index]
                if ref_para != p:
                    ref_para._element.addnext(p._element)

            return {
                "success": True,
                "message": f"审阅标记 #{review_id} 插入成功",
                "review_id": review_id,
                "comment_type": comment_type,
                "marker_text": marker_text,
                "note": "跨平台模式下以格式化段落标记，非 Word 原生批注",
            }
        except Exception as e:
            return {"success": False, "error": f"插入审阅标记失败: {str(e)}"}

    def highlight_text(
        self,
        text_to_highlight: str,
        color: str = "yellow",
    ) -> Dict[str, Any]:
        """
        对匹配文本加高亮或显眼样式。

        :param text_to_highlight: 要高亮的文本
        :param color: 高亮颜色 (yellow/red/green/blue/gray)
        :return: 高亮结果
        """
        if not DOCX_AVAILABLE:
            return {"success": False, "error": "python-docx 库不可用"}

        try:
            if not self.document:
                return {"success": False, "error": "文档未加载"}

            from docx.enum.text import WD_COLOR_INDEX
            color_index_map = {
                "yellow": WD_COLOR_INDEX.YELLOW,
                "red": WD_COLOR_INDEX.RED,
                "green": WD_COLOR_INDEX.GREEN,
                "blue": WD_COLOR_INDEX.BLUE,
                "gray": WD_COLOR_INDEX.GRAY_25,
            }
            wd_color = color_index_map.get(color, WD_COLOR_INDEX.YELLOW)

            modifications = 0
            for para in self.document.paragraphs:
                for run in para.runs:
                    if text_to_highlight in run.text:
                        run.font.highlight_color = wd_color
                        modifications += 1

            for table in self.document.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for para in cell.paragraphs:
                            for run in para.runs:
                                if text_to_highlight in run.text:
                                    run.font.highlight_color = wd_color
                                    modifications += 1

            return {
                "success": True,
                "message": f"高亮完成，共标记 {modifications} 处文本",
                "modifications": modifications,
                "color": color,
            }
        except Exception as e:
            return {"success": False, "error": f"高亮文本失败: {str(e)}"}

    def mark_paragraph_as_todo(
        self,
        paragraph_index: int,
        todo_note: str = "",
    ) -> Dict[str, Any]:
        """
        给某段增加 TODO 标记（在段落前插入格式化 TODO 行）。

        :param paragraph_index: 段落索引
        :param todo_note: TODO 备注
        :return: 标记结果
        """
        if not DOCX_AVAILABLE:
            return {"success": False, "error": "python-docx 库不可用"}

        try:
            if not self.document:
                return {"success": False, "error": "文档未加载"}

            if paragraph_index < 0 or paragraph_index >= len(self.document.paragraphs):
                return {"success": False, "error": f"段落索引 {paragraph_index} 超出范围"}

            self._todo_counter += 1
            para = self.document.paragraphs[paragraph_index]

            todo_text = f"[TODO#{self._todo_counter}]"
            if todo_note:
                todo_text += f" {todo_note}"

            todo_run = para.insert_paragraph_before("")
            run = todo_run.add_run(todo_text)
            run.bold = True
            run.font.color.rgb = RGBColor(255, 140, 0)
            run.font.size = Pt(9)

            return {
                "success": True,
                "message": f"段落 {paragraph_index} 已标记为 TODO #{self._todo_counter}",
                "todo_id": self._todo_counter,
                "todo_text": todo_text,
                "note": "跨平台模式下以格式化段落标记",
            }
        except Exception as e:
            return {"success": False, "error": f"标记 TODO 失败: {str(e)}"}

    def review_section(
        self,
        section_heading: str,
        checks: Optional[List[str]] = None,
    ) -> Dict[str, Any]:
        """
        读取该 section 文本并生成结构化审阅结果。
        不直接修改文档，返回可供 Xenon 后续处理的结果。

        :param section_heading: 章节标题文本
        :param checks: 检查项列表 ["citation", "structure", "clarity", "repetition"]
        :return: 结构化审阅结果
        """
        if not DOCX_AVAILABLE:
            return {"success": False, "error": "python-docx 库不可用"}

        try:
            if not self.document:
                return {"success": False, "error": "文档未加载"}

            if checks is None:
                checks = ["citation", "structure", "clarity"]

            section_start = None
            section_level = None
            for i, para in enumerate(self.document.paragraphs):
                if para.style.name.startswith("Heading") and section_heading in para.text:
                    section_start = i
                    try:
                        section_level = int(para.style.name.split()[-1])
                    except (ValueError, IndexError):
                        section_level = 1
                    break

            if section_start is None:
                return {"success": False, "error": f"未找到章节: {section_heading}"}

            section_end = len(self.document.paragraphs)
            for i in range(section_start + 1, len(self.document.paragraphs)):
                para = self.document.paragraphs[i]
                if para.style.name.startswith("Heading"):
                    try:
                        plevel = int(para.style.name.split()[-1])
                    except (ValueError, IndexError):
                        plevel = 9
                    if plevel <= section_level:
                        section_end = i
                        break

            paragraphs_text = []
            for i in range(section_start, section_end):
                paragraphs_text.append(self.document.paragraphs[i].text)

            full_text = "\n".join(paragraphs_text)
            word_count = sum(len(p) for p in paragraphs_text)
            para_count = section_end - section_start

            findings = []

            if "citation" in checks:
                citation_patterns = [r'\[\d+\]', r'\(\w+,\s*\d{4}\)', r'\[\w+\d{4}\]']
                has_citation = any(re.search(pat, full_text) for pat in citation_patterns)
                if not has_citation and para_count > 3:
                    findings.append({
                        "type": "citation",
                        "severity": "warning",
                        "message": f"章节 '{section_heading}' 未发现引用标记，建议补充参考文献引用",
                    })

            if "structure" in checks:
                sub_headings = []
                for i in range(section_start + 1, section_end):
                    p = self.document.paragraphs[i]
                    if p.style.name.startswith("Heading"):
                        sub_headings.append(p.text)
                if para_count > 10 and not sub_headings:
                    findings.append({
                        "type": "structure",
                        "severity": "info",
                        "message": f"章节 '{section_heading}' 有 {para_count} 个段落但无子标题，考虑添加子节",
                    })

            if "clarity" in checks:
                for i, pt in enumerate(paragraphs_text):
                    if len(pt) > 500:
                        findings.append({
                            "type": "clarity",
                            "severity": "info",
                            "message": f"段落 {i} 过长（{len(pt)} 字符），建议拆分",
                        })

            if "repetition" in checks:
                for i in range(1, len(paragraphs_text)):
                    if paragraphs_text[i] and paragraphs_text[i - 1]:
                        prefix_len = min(10, len(paragraphs_text[i]), len(paragraphs_text[i - 1]))
                        if paragraphs_text[i][:prefix_len] == paragraphs_text[i - 1][:prefix_len]:
                            findings.append({
                                "type": "repetition",
                                "severity": "info",
                                "message": f"段落 {i} 与前一段开头相似，可能存在重复",
                            })

            return {
                "success": True,
                "section_heading": section_heading,
                "section_level": section_level,
                "paragraph_count": para_count,
                "word_count": word_count,
                "checks_performed": checks,
                "findings": findings,
                "findings_count": len(findings),
                "note": "此为结构化审阅结果，未直接修改文档",
            }
        except Exception as e:
            return {"success": False, "error": f"审阅章节失败: {str(e)}"}


class WordToolManager:
    def __init__(self):
        """
        Word 工具管理器
        """
        self.handlers = {}
    
    def load_document(self, file_path: str) -> Dict[str, Any]:
        """
        加载 Word 文档
        :param file_path: Word 文件路径
        """
        try:
            handler = WordHandler(file_path)
            result = handler.load_document()
            if result["success"]:
                self.handlers[file_path] = handler
            return result
        except Exception as e:
            return {"success": False, "error": f"加载Word文档失败: {str(e)}"}
    
    def create_new_document(self, file_path: str = None) -> Dict[str, Any]:
        """
        创建新的空白文档
        :param file_path: Word 文件路径（可选）
        """
        try:
            handler = WordHandler(file_path)
            result = handler.create_new_document()
            if result["success"]:
                # 使用一个临时的唯一键来存储新文档
                temp_key = f"new_doc_{id(handler)}"
                self.handlers[temp_key] = handler
                result["doc_key"] = temp_key
            return result
        except Exception as e:
            return {"success": False, "error": f"创建新文档失败: {str(e)}"}
    
    def read_paragraphs(self, file_path: str, start: int = 0, end: Optional[int] = None) -> Dict[str, Any]:
        """
        读取文档段落内容
        :param file_path: Word 文件路径
        :param start: 起始段落索引
        :param end: 结束段落索引（不包含）
        """
        try:
            if file_path not in self.handlers:
                handler = WordHandler(file_path)
                load_result = handler.load_document()
                if not load_result["success"]:
                    return load_result
                self.handlers[file_path] = handler
            
            return self.handlers[file_path].read_paragraphs(start, end)
        except Exception as e:
            return {"success": False, "error": f"读取段落失败: {str(e)}"}
    
    def add_paragraph(self, file_path: str, text: str, style: str = "Normal", position: int = -1) -> Dict[str, Any]:
        """
        添加段落到文档
        :param file_path: Word 文件路径
        :param text: 要添加的文本
        :param style: 段落样式
        :param position: 插入位置，-1表示末尾
        """
        try:
            if file_path not in self.handlers:
                handler = WordHandler(file_path)
                load_result = handler.load_document()
                if not load_result["success"]:
                    return load_result
                self.handlers[file_path] = handler
            
            return self.handlers[file_path].add_paragraph(text, style, position)
        except Exception as e:
            return {"success": False, "error": f"添加段落失败: {str(e)}"}
    
    def add_heading(self, file_path: str, text: str, level: int = 1, position: int = -1) -> Dict[str, Any]:
        """
        添加标题
        :param file_path: Word 文件路径
        :param text: 标题文本
        :param level: 标题级别（1-9）
        :param position: 插入位置，-1表示末尾
        """
        try:
            if file_path not in self.handlers:
                handler = WordHandler(file_path)
                load_result = handler.load_document()
                if not load_result["success"]:
                    return load_result
                self.handlers[file_path] = handler
            
            return self.handlers[file_path].add_heading(text, level, position)
        except Exception as e:
            return {"success": False, "error": f"添加标题失败: {str(e)}"}
    
    def replace_text(self, file_path: str, old_text: str, new_text: str, replace_all: bool = True, auto_save: bool = True) -> Dict[str, Any]:
        """
        替换文档中的文本
        :param file_path: Word 文件路径
        :param old_text: 要替换的文本
        :param new_text: 新文本
        :param replace_all: 是否替换所有匹配项
        :param auto_save: 是否在替换成功后立即保存到原文件，默认 True
        """
        try:
            if file_path not in self.handlers:
                handler = WordHandler(file_path)
                load_result = handler.load_document()
                if not load_result["success"]:
                    return load_result
                self.handlers[file_path] = handler
            
            result = self.handlers[file_path].replace_text(old_text, new_text, replace_all)
            if result.get("success") and result.get("replacements", 0) > 0:
                if auto_save:
                    save_result = self.handlers[file_path].save_document(file_path)
                    if not save_result["success"]:
                        return {
                            "success": False,
                            "error": f"文本已在内存中替换，但保存失败: {save_result.get('error')}",
                            "replace_result": result,
                        }
                    result["saved"] = True
                    result["save_result"] = save_result
                else:
                    result["saved"] = False
                    result["note"] = "文本已在内存中替换；请调用 save_document 保存到磁盘"
            return result
        except Exception as e:
            return {"success": False, "error": f"替换文本失败: {str(e)}"}
    
    def create_table(self, file_path: str, rows: int, cols: int, headers: Optional[List[str]] = None, data: Optional[List[List[str]]] = None, position: int = -1) -> Dict[str, Any]:
        """
        在文档中创建表格
        :param file_path: Word 文件路径
        :param rows: 表格行数
        :param cols: 表格列数
        :param headers: 表头列表（可选）
        :param data: 表格数据，二维列表格式（可选）
        :param position: 插入位置，-1表示末尾
        """
        try:
            if file_path not in self.handlers:
                handler = WordHandler(file_path)
                load_result = handler.load_document()
                if not load_result["success"]:
                    return load_result
                self.handlers[file_path] = handler
            
            return self.handlers[file_path].create_table(rows, cols, headers, data, position)
        except Exception as e:
            return {"success": False, "error": f"创建表格失败: {str(e)}"}
    
    def add_table_from_data(self, file_path: str, data_list: List[List[str]], headers: Optional[List[str]] = None, position: int = -1) -> Dict[str, Any]:
        """
        从二维数据列表直接创建表格
        :param file_path: Word 文件路径
        :param data_list: 二维数据列表
        :param headers: 表头列表（可选）
        :param position: 插入位置，-1表示末尾
        """
        try:
            if file_path not in self.handlers:
                handler = WordHandler(file_path)
                load_result = handler.load_document()
                if not load_result["success"]:
                    return load_result
                self.handlers[file_path] = handler
            
            return self.handlers[file_path].add_table_from_data(data_list, headers, position)
        except Exception as e:
            return {"success": False, "error": f"从数据创建表格失败: {str(e)}"}
    
    def update_table_cell(self, file_path: str, table_index: int, row: int, col: int, value: str) -> Dict[str, Any]:
        """
        更新指定表格的单元格内容
        :param file_path: Word 文件路径
        :param table_index: 表格索引
        :param row: 行索引
        :param col: 列索引
        :param value: 新的单元格值
        """
        try:
            if file_path not in self.handlers:
                handler = WordHandler(file_path)
                load_result = handler.load_document()
                if not load_result["success"]:
                    return load_result
                self.handlers[file_path] = handler
            
            return self.handlers[file_path].update_table_cell(table_index, row, col, value)
        except Exception as e:
            return {"success": False, "error": f"更新表格单元格失败: {str(e)}"}
    
    def get_table_data(self, file_path: str, table_index: int) -> Dict[str, Any]:
        """
        获取指定表格的所有数据
        :param file_path: Word 文件路径
        :param table_index: 表格索引
        """
        try:
            if file_path not in self.handlers:
                handler = WordHandler(file_path)
                load_result = handler.load_document()
                if not load_result["success"]:
                    return load_result
                self.handlers[file_path] = handler
            
            return self.handlers[file_path].get_table_data(table_index)
        except Exception as e:
            return {"success": False, "error": f"获取表格数据失败: {str(e)}"}
    
    def set_paragraph_format(self, file_path: str, paragraph_index: int, alignment: str = None, spacing_before: float = None, spacing_after: float = None) -> Dict[str, Any]:
        """
        设置段落格式
        :param file_path: Word 文件路径
        :param paragraph_index: 段落索引
        :param alignment: 对齐方式 ('left', 'center', 'right', 'justify')
        :param spacing_before: 段前间距（磅）
        :param spacing_after: 段后间距（磅）
        """
        try:
            if file_path not in self.handlers:
                handler = WordHandler(file_path)
                load_result = handler.load_document()
                if not load_result["success"]:
                    return load_result
                self.handlers[file_path] = handler
            
            return self.handlers[file_path].set_paragraph_format(paragraph_index, alignment, spacing_before, spacing_after)
        except Exception as e:
            return {"success": False, "error": f"设置段落格式失败: {str(e)}"}
    
    def set_font_style(self, file_path: str, text_range: str, bold: bool = False, italic: bool = False, underline: bool = False, font_size: float = None, font_name: str = None) -> Dict[str, Any]:
        """
        设置字体样式
        :param file_path: Word 文件路径
        :param text_range: 文本范围
        :param bold: 是否加粗
        :param italic: 是否斜体
        :param underline: 是否下划线
        :param font_size: 字体大小（磅）
        :param font_name: 字体名称
        """
        try:
            if file_path not in self.handlers:
                handler = WordHandler(file_path)
                load_result = handler.load_document()
                if not load_result["success"]:
                    return load_result
                self.handlers[file_path] = handler
            
            return self.handlers[file_path].set_font_style(text_range, bold, italic, underline, font_size, font_name)
        except Exception as e:
            return {"success": False, "error": f"设置字体样式失败: {str(e)}"}
    
    def insert_page_break(self, file_path: str, position: int = -1) -> Dict[str, Any]:
        """
        插入分页符
        :param file_path: Word 文件路径
        :param position: 插入位置，-1表示末尾
        """
        try:
            if file_path not in self.handlers:
                handler = WordHandler(file_path)
                load_result = handler.load_document()
                if not load_result["success"]:
                    return load_result
                self.handlers[file_path] = handler
            
            return self.handlers[file_path].insert_page_break(position)
        except Exception as e:
            return {"success": False, "error": f"插入分页符失败: {str(e)}"}
    
    def add_section_break(self, file_path: str, position: int = -1) -> Dict[str, Any]:
        """
        插入分节符
        :param file_path: Word 文件路径
        :param position: 插入位置，-1表示末尾
        """
        try:
            if file_path not in self.handlers:
                handler = WordHandler(file_path)
                load_result = handler.load_document()
                if not load_result["success"]:
                    return load_result
                self.handlers[file_path] = handler
            
            return self.handlers[file_path].add_section_break(position)
        except Exception as e:
            return {"success": False, "error": f"插入分节符失败: {str(e)}"}
    
    def get_document_structure(self, file_path: str) -> Dict[str, Any]:
        """
        获取文档结构信息
        :param file_path: Word 文件路径
        """
        try:
            if file_path not in self.handlers:
                handler = WordHandler(file_path)
                load_result = handler.load_document()
                if not load_result["success"]:
                    return load_result
                self.handlers[file_path] = handler
            
            return self.handlers[file_path].get_document_structure()
        except Exception as e:
            return {"success": False, "error": f"获取文档结构失败: {str(e)}"}
    
    def import_excel_to_tables(self, file_path: str, excel_file_path: str, sheet_names: Optional[List[str]] = None) -> Dict[str, Any]:
        """
        将Excel工作表直接转换为Word表格
        :param file_path: Word 文件路径
        :param excel_file_path: Excel文件路径
        :param sheet_names: 工作表名称列表（可选）
        """
        try:
            if file_path not in self.handlers:
                handler = WordHandler(file_path)
                load_result = handler.load_document()
                if not load_result["success"]:
                    return load_result
                self.handlers[file_path] = handler
            
            return self.handlers[file_path].import_excel_to_tables(excel_file_path, sheet_names)
        except Exception as e:
            return {"success": False, "error": f"导入Excel到表格失败: {str(e)}"}
    
    def export_tables_to_excel(self, file_path: str, output_path: str) -> Dict[str, Any]:
        """
        将文档中的表格导出到Excel
        :param file_path: Word 文件路径
        :param output_path: 输出Excel文件路径
        """
        try:
            if file_path not in self.handlers:
                handler = WordHandler(file_path)
                load_result = handler.load_document()
                if not load_result["success"]:
                    return load_result
                self.handlers[file_path] = handler
            
            return self.handlers[file_path].export_tables_to_excel(output_path)
        except Exception as e:
            return {"success": False, "error": f"导出表格到Excel失败: {str(e)}"}
    
    def save_document(self, file_path: str, output_path: Optional[str] = None) -> Dict[str, Any]:
        """
        保存文档
        :param file_path: Word 文件路径
        :param output_path: 输出路径（可选，默认使用原路径）
        """
        try:
            if file_path not in self.handlers:
                handler = WordHandler(file_path)
                load_result = handler.load_document()
                if not load_result["success"]:
                    return load_result
                self.handlers[file_path] = handler
            
            return self.handlers[file_path].save_document(output_path)
        except Exception as e:
            return {"success": False, "error": f"保存文档失败: {str(e)}"}

    def merge_table_cells(self, file_path: str, table_index: int, top_row: int, left_col: int, bottom_row: int, right_col: int) -> Dict[str, Any]:
        """
        合并表格中的单元格
        :param file_path: Word 文件路径
        :param table_index: 表格索引
        :param top_row: 顶部行索引
        :param left_col: 左侧列索引
        :param bottom_row: 底部行索引
        :param right_col: 右侧列索引
        """
        try:
            if file_path not in self.handlers:
                handler = WordHandler(file_path)
                load_result = handler.load_document()
                if not load_result["success"]:
                    return load_result
                self.handlers[file_path] = handler
            
            return self.handlers[file_path].merge_table_cells(table_index, top_row, left_col, bottom_row, right_col)
        except Exception as e:
            return {"success": False, "error": f"合并单元格失败: {str(e)}"}

    def split_table_cell(self, file_path: str, table_index: int, row: int, col: int, v_merge: bool = True) -> Dict[str, Any]:
        """
        拆分表格中的已合并单元格。
        注意：当前实现会明确返回“暂不支持可靠拆分”，用于避免直接改 XML 导致文档损坏。
        :param file_path: Word 文件路径
        :param table_index: 表格索引
        :param row: 行索引
        :param col: 列索引
        :param v_merge: 保留参数，当前不生效
        """
        try:
            if file_path not in self.handlers:
                handler = WordHandler(file_path)
                load_result = handler.load_document()
                if not load_result["success"]:
                    return load_result
                self.handlers[file_path] = handler
            
            return self.handlers[file_path].split_table_cell(table_index, row, col, v_merge)
        except Exception as e:
            return {"success": False, "error": f"拆分单元格失败: {str(e)}"}

    def adjust_cell_width_height(self, file_path: str, table_index: int, row: int, col: int, width_inches: float = None, height_inches: float = None) -> Dict[str, Any]:
        """
        调整表格单元格的宽度和高度
        :param file_path: Word 文件路径
        :param table_index: 表格索引
        :param row: 行索引
        :param col: 列索引
        :param width_inches: 宽度（英寸）
        :param height_inches: 高度（英寸）
        """
        try:
            if file_path not in self.handlers:
                handler = WordHandler(file_path)
                load_result = handler.load_document()
                if not load_result["success"]:
                    return load_result
                self.handlers[file_path] = handler
            
            return self.handlers[file_path].adjust_cell_width_height(table_index, row, col, width_inches, height_inches)
        except Exception as e:
            return {"success": False, "error": f"调整单元格尺寸失败: {str(e)}"}

    def insert_picture_with_wrap(self, word_path: str, image_path: str, wrap_type: str = "嵌入型", 
                                   position_x: Optional[float] = None, position_y: Optional[float] = None,
                                   width: Optional[float] = None, height: Optional[float] = None,
                                   target_text: Optional[str] = None) -> Dict[str, Any]:
        """
        在Word文档中插入图片并设置文字环绕类型。
        跨平台规则：
        1. "嵌入型"(I) 为跨平台能力，基于 python-docx，可在 Windows / Linux / macOS 使用。
        2. 其他环绕类型依赖 Microsoft Word COM 自动化，只能在 Windows + Word + pywin32 环境下使用。
        3. 如果运行环境不是 Windows，调用本工具时应优先选择 "嵌入型"(I)。
        
        :param word_path: Word文档路径
        :param image_path: 图片路径
        :param wrap_type: 文字环绕类型，可选值：
                         - "四周型"(Q) - wdWrapSquare = 0
                         - "紧密型"(T) - wdWrapTight = 1
                         - "穿越型"(H) - wdWrapThrough = 2
                         - "上下型"(O) - wdWrapTopBottom = 3
                         - "衬于文字下方"(B) - wdWrapBehind = 5
                         - "浮于文字上方"(F) - wdWrapFront = 4
                         - "嵌入型"(I) - wdWrapInline = 7
        :param position_x: 图片水平位置（磅，1英寸=72磅），仅对 Windows 下的非嵌入型有效
        :param position_y: 图片垂直位置（磅，1英寸=72磅），仅对 Windows 下的非嵌入型有效
        :param width: 图片宽度（磅），None表示保持原始比例
        :param height: 图片高度（磅），None表示保持原始比例
        :param target_text: 目标文本位置，如果提供则在该文本处插入图片；跨平台模式下仅支持插入嵌入型图片
        :return: 包含成功状态和结果/错误的字典
        """
        try:
            handler = WordHandler()
            return handler.insert_picture_with_wrap(word_path, image_path, wrap_type, position_x, position_y, width, height, target_text)
        except Exception as e:
            return {"success": False, "error": f"插入图片失败: {str(e)}"}

    # ════════════════════════════════════════════════════════════
    # 学术论文能力（对外暴露）
    # ════════════════════════════════════════════════════════════

    def _get_handler(self, file_path: str) -> Tuple[Optional['WordHandler'], Optional[Dict[str, Any]]]:
        """获取或创建 handler 的辅助方法。返回 (handler, error_dict)，handler 为 None 时 error_dict 不为 None。"""
        if file_path not in self.handlers:
            handler = WordHandler(file_path)
            load_result = handler.load_document()
            if not load_result["success"]:
                return None, load_result
            self.handlers[file_path] = handler
        return self.handlers[file_path], None

    def initialize_academic_template(
        self,
        file_path: str,
        title: str = "",
        author: str = "",
        institution: str = "",
        date_text: str = "",
        language: str = "zh",
        paper_type: str = "article",
        style_preset: str = "cn_academic",
        include_cover: bool = True,
        include_abstract_placeholders: bool = True,
        include_toc_placeholder: bool = True,
    ) -> Dict[str, Any]:
        """
        一键创建适合论文写作的文档基线。

        :param file_path: Word 文件路径
        :param title: 论文标题
        :param author: 作者
        :param institution: 机构
        :param date_text: 日期文本
        :param language: "zh" / "en"
        :param paper_type: "article" / "thesis" / "report"
        :param style_preset: "cn_academic" / "ieee_like" / "apa_like"
        :param include_cover: 是否生成封面
        :param include_abstract_placeholders: 是否生成摘要占位章节
        :param include_toc_placeholder: 是否插入目录占位
        """
        try:
            handler = WordHandler(file_path)
            result = handler.initialize_academic_template(
                title=title, author=author, institution=institution,
                date_text=date_text, language=language, paper_type=paper_type,
                style_preset=style_preset, include_cover=include_cover,
                include_abstract_placeholders=include_abstract_placeholders,
                include_toc_placeholder=include_toc_placeholder,
            )
            if result["success"]:
                self.handlers[file_path] = handler
            return result
        except Exception as e:
            return {"success": False, "error": f"初始化学术模板失败: {str(e)}"}

    def insert_caption(
        self, file_path: str, target_type: str = "figure", caption_text: str = "",
        numbering_scope: str = "global", position: str = "after",
        target_index: Optional[int] = None,
    ) -> Dict[str, Any]:
        """
        插入题注（图/表/公式），自动生成编号。

        :param file_path: Word 文件路径
        :param target_type: "figure" / "table" / "equation"
        :param caption_text: 题注文字
        :param numbering_scope: 编号范围
        :param position: "before" / "after"
        :param target_index: 目标段落索引
        """
        try:
            handler, err = self._get_handler(file_path)
            if err:
                return err
            return handler.insert_caption(target_type, caption_text, numbering_scope, position, target_index)
        except Exception as e:
            return {"success": False, "error": f"插入题注失败: {str(e)}"}

    def insert_table_with_caption(
        self, file_path: str, rows: int, cols: int, caption_text: str = "",
        headers: Optional[List[str]] = None, data: Optional[List[List[str]]] = None,
        position: int = -1,
    ) -> Dict[str, Any]:
        """
        封装"建表 + 题注"。

        :param file_path: Word 文件路径
        :param rows: 行数
        :param cols: 列数
        :param caption_text: 题注文字
        :param headers: 表头列表
        :param data: 数据二维列表
        :param position: 插入位置
        """
        try:
            handler, err = self._get_handler(file_path)
            if err:
                return err
            return handler.insert_table_with_caption(rows, cols, caption_text, headers, data, position)
        except Exception as e:
            return {"success": False, "error": f"插入带题注表格失败: {str(e)}"}

    def insert_picture_with_caption(
        self, word_path: str, image_path: str, caption_text: str = "",
        wrap_type: str = "嵌入型", width: Optional[float] = None,
        height: Optional[float] = None, target_text: Optional[str] = None,
    ) -> Dict[str, Any]:
        """
        封装"插图 + 题注"。

        :param word_path: Word 文件路径
        :param image_path: 图片路径
        :param caption_text: 题注文字
        :param wrap_type: 文字环绕类型
        :param width: 宽度
        :param height: 高度
        :param target_text: 目标文本
        """
        try:
            handler, err = self._get_handler(word_path)
            if err:
                return err
            return handler.insert_picture_with_caption(
                word_path, image_path, caption_text, wrap_type, width, height, target_text
            )
        except Exception as e:
            return {"success": False, "error": f"插入带题注图片失败: {str(e)}"}

    def insert_toc_placeholder(self, file_path: str, language: str = "zh") -> Dict[str, Any]:
        """
        插入目录占位区。

        :param file_path: Word 文件路径
        :param language: "zh" / "en"
        """
        try:
            handler, err = self._get_handler(file_path)
            if err:
                return err
            return handler.insert_toc_placeholder(language)
        except Exception as e:
            return {"success": False, "error": f"插入目录占位失败: {str(e)}"}

    def insert_list_of_figures_placeholder(self, file_path: str, language: str = "zh") -> Dict[str, Any]:
        """
        插入图目录占位区。

        :param file_path: Word 文件路径
        :param language: "zh" / "en"
        """
        try:
            handler, err = self._get_handler(file_path)
            if err:
                return err
            return handler.insert_list_of_figures_placeholder(language)
        except Exception as e:
            return {"success": False, "error": f"插入图目录占位失败: {str(e)}"}

    def insert_list_of_tables_placeholder(self, file_path: str, language: str = "zh") -> Dict[str, Any]:
        """
        插入表目录占位区。

        :param file_path: Word 文件路径
        :param language: "zh" / "en"
        """
        try:
            handler, err = self._get_handler(file_path)
            if err:
                return err
            return handler.insert_list_of_tables_placeholder(language)
        except Exception as e:
            return {"success": False, "error": f"插入表目录占位失败: {str(e)}"}

    def add_reference_entry(
        self, file_path: str, cite_key: str, authors: str = "", title: str = "",
        year: str = "", journal: str = "", volume: str = "", issue: str = "",
        pages: str = "", publisher: str = "", doi: str = "", url: str = "",
        entry_type: str = "article", style: str = "gbt7714",
    ) -> Dict[str, Any]:
        """
        登记一条参考文献条目。

        :param file_path: Word 文件路径
        :param cite_key: 引用键
        :param authors: 作者
        :param title: 标题
        :param year: 年份
        :param journal: 期刊
        :param volume: 卷
        :param issue: 期
        :param pages: 页码
        :param publisher: 出版社
        :param doi: DOI
        :param url: URL
        :param entry_type: article / book / conference / webpage / thesis / misc
        :param style: gbt7714 / ieee / apa_like
        """
        try:
            handler, err = self._get_handler(file_path)
            if err:
                return err
            return handler.add_reference_entry(
                cite_key, authors, title, year, journal, volume, issue,
                pages, publisher, doi, url, entry_type, style
            )
        except Exception as e:
            return {"success": False, "error": f"登记参考文献失败: {str(e)}"}

    def insert_citation(
        self, file_path: str, cite_keys: List[str], style: str = "gbt7714",
        position_mode: str = "append_current_paragraph",
        target_text: Optional[str] = None,
    ) -> Dict[str, Any]:
        """
        在正文中插入轻量引用标记。

        :param file_path: Word 文件路径
        :param cite_keys: 引用键列表
        :param style: "gbt7714" / "ieee" / "apa_like"
        :param position_mode: "append_current_paragraph" / "insert_new_paragraph" / "target_text"
        :param target_text: 目标文本
        """
        try:
            handler, err = self._get_handler(file_path)
            if err:
                return err
            return handler.insert_citation(cite_keys, style, position_mode, target_text)
        except Exception as e:
            return {"success": False, "error": f"插入引用失败: {str(e)}"}

    def generate_bibliography(
        self, file_path: str, style: str = "gbt7714",
        heading_text: str = "", sort_mode: str = "appearance",
    ) -> Dict[str, Any]:
        """
        自动生成参考文献列表。

        :param file_path: Word 文件路径
        :param style: "gbt7714" / "ieee" / "apa_like"
        :param heading_text: 标题
        :param sort_mode: "appearance" / "alphabetical"
        """
        try:
            handler, err = self._get_handler(file_path)
            if err:
                return err
            return handler.generate_bibliography(style, heading_text, sort_mode)
        except Exception as e:
            return {"success": False, "error": f"生成参考文献失败: {str(e)}"}

    def insert_heading_after_text(
        self, file_path: str, text_anchor: str, heading_text: str, level: int = 1,
    ) -> Dict[str, Any]:
        """
        在指定文本锚点后插入标题。

        :param file_path: Word 文件路径
        :param text_anchor: 锚点文本
        :param heading_text: 标题文本
        :param level: 标题级别
        """
        try:
            handler, err = self._get_handler(file_path)
            if err:
                return err
            return handler.insert_heading_after_text(text_anchor, heading_text, level)
        except Exception as e:
            return {"success": False, "error": f"插入标题失败: {str(e)}"}

    def insert_section_block(
        self, file_path: str, heading_text: str, level: int = 1,
        paragraphs: Optional[List[str]] = None,
        after_heading: Optional[str] = None, after_text: Optional[str] = None,
        page_break_before: bool = False,
    ) -> Dict[str, Any]:
        """
        一次插入一整节。

        :param file_path: Word 文件路径
        :param heading_text: 标题文本
        :param level: 标题级别
        :param paragraphs: 段落列表
        :param after_heading: 在此标题后插入
        :param after_text: 在此文本后插入
        :param page_break_before: 节前分页
        """
        try:
            handler, err = self._get_handler(file_path)
            if err:
                return err
            return handler.insert_section_block(
                heading_text, level, paragraphs, after_heading, after_text, page_break_before
            )
        except Exception as e:
            return {"success": False, "error": f"插入节块失败: {str(e)}"}

    def move_section(
        self, file_path: str, section_heading: str, target_after_heading: str,
    ) -> Dict[str, Any]:
        """
        章节块移动。

        :param file_path: Word 文件路径
        :param section_heading: 源章节标题
        :param target_after_heading: 目标位置标题
        """
        try:
            handler, err = self._get_handler(file_path)
            if err:
                return err
            return handler.move_section(section_heading, target_after_heading)
        except Exception as e:
            return {"success": False, "error": f"移动章节失败: {str(e)}"}

    def get_document_outline(self, file_path: str) -> Dict[str, Any]:
        """
        获取文档大纲（标题层级）。

        :param file_path: Word 文件路径
        """
        try:
            handler, err = self._get_handler(file_path)
            if err:
                return err
            return handler.get_document_outline()
        except Exception as e:
            return {"success": False, "error": f"获取文档大纲失败: {str(e)}"}

    def insert_review_comment_block(
        self, file_path: str, comment_text: str, comment_type: str = "todo",
        target_text: Optional[str] = None, after_paragraph_index: Optional[int] = None,
    ) -> Dict[str, Any]:
        """
        插入格式化审阅块。

        :param file_path: Word 文件路径
        :param comment_text: 审阅内容
        :param comment_type: todo / citation_needed / rewrite / logic_check / format_check
        :param target_text: 锚点文本
        :param after_paragraph_index: 段落索引
        """
        try:
            handler, err = self._get_handler(file_path)
            if err:
                return err
            return handler.insert_review_comment_block(
                comment_text, comment_type, target_text, after_paragraph_index
            )
        except Exception as e:
            return {"success": False, "error": f"插入审阅标记失败: {str(e)}"}

    def highlight_text(
        self, file_path: str, text_to_highlight: str, color: str = "yellow",
    ) -> Dict[str, Any]:
        """
        对匹配文本加高亮。

        :param file_path: Word 文件路径
        :param text_to_highlight: 要高亮的文本
        :param color: 颜色 (yellow/red/green/blue/gray)
        """
        try:
            handler, err = self._get_handler(file_path)
            if err:
                return err
            return handler.highlight_text(text_to_highlight, color)
        except Exception as e:
            return {"success": False, "error": f"高亮文本失败: {str(e)}"}

    def mark_paragraph_as_todo(
        self, file_path: str, paragraph_index: int, todo_note: str = "",
    ) -> Dict[str, Any]:
        """
        给某段增加 TODO 标记。

        :param file_path: Word 文件路径
        :param paragraph_index: 段落索引
        :param todo_note: TODO 备注
        """
        try:
            handler, err = self._get_handler(file_path)
            if err:
                return err
            return handler.mark_paragraph_as_todo(paragraph_index, todo_note)
        except Exception as e:
            return {"success": False, "error": f"标记 TODO 失败: {str(e)}"}

    def review_section(
        self, file_path: str, section_heading: str,
        checks: Optional[List[str]] = None,
    ) -> Dict[str, Any]:
        """
        结构化审阅章节。

        :param file_path: Word 文件路径
        :param section_heading: 章节标题
        :param checks: 检查项列表
        """
        try:
            handler, err = self._get_handler(file_path)
            if err:
                return err
            return handler.review_section(section_heading, checks)
        except Exception as e:
            return {"success": False, "error": f"审阅章节失败: {str(e)}"}


# 为兼容性提供模块级函数
def create_word_tool_manager():
    """
    创建 Word 工具管理器实例
    """
    return WordToolManager()


def example_usage():
    """
    使用示例：展示如何使用 WordHandler 创建文档并添加内容
    """
    # 创建 WordHandler 实例
    word_handler = WordHandler()
    
    # 创建新文档
    result = word_handler.create_new_document()
    print(f"创建文档结果: {result}")
    
    # 添加标题
    result = word_handler.add_heading("这是一个标题", level=1)
    print(f"添加标题结果: {result}")
    
    # 添加段落
    result = word_handler.add_paragraph("这是第一个段落的内容。")
    print(f"添加段落结果: {result}")
    
    # 创建表格
    headers = ["姓名", "年龄", "城市"]
    data = [
        ["张三", "25", "北京"],
        ["李四", "30", "上海"],
        ["王五", "28", "广州"]
    ]
    result = word_handler.create_table(3, 3, headers=headers, data=data)
    print(f"创建表格结果: {result}")
    
    # 保存文档
    result = word_handler.save_document("./example.docx")
    print(f"保存文档结果: {result}")


def batch_replace_texts(file_path: str, replacements: Dict[str, str]) -> Dict[str, Any]:
    """
    批量替换文本功能
    :param file_path: Word文档路径
    :param replacements: 要替换的文本映射字典，格式为 {旧文本: 新文本}
    :return: 包含操作结果信息的字典
    """
    try:
        handler = WordHandler(file_path)
        load_result = handler.load_document()
        if not load_result["success"]:
            return {"success": False, "error": f"加载文档失败: {load_result['error']}"}
        
        total_replacements = 0
        for old_text, new_text in replacements.items():
            result = handler.replace_text(old_text, new_text, replace_all=True)
            if result["success"]:
                total_replacements += result.get("replacements", 0)
        
        save_result = handler.save_document(file_path)
        if not save_result["success"]:
            return {"success": False, "error": f"保存文档失败: {save_result['error']}"}
        
        return {
            "success": True,
            "message": f"批量替换完成，总共替换了 {total_replacements} 处文本",
            "total_replacements": total_replacements
        }
    except Exception as e:
        return {"success": False, "error": f"批量替换失败: {str(e)}"}
